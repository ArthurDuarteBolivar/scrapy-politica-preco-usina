import requests
import mysql.connector
import unidecode
import scrapy
import requests
from docx import Document
import pandas

start_row = 20  
end_row = 33
num_rows = end_row - start_row


df = pandas.read_excel("GESTÃO DE AÇÕES E-COMMERCE.xlsx", usecols='C:K', skiprows=start_row, nrows=num_rows, engine='openpyxl')

df.columns = ['PRODUTO', 'SITE', 'COLUNA3', 'CLÁSSICO ML', 'COLUNA5', 'PREMIUM ML', 'COLUNA7', 'MARKETPLACES', 'COLUNA9']


fonteBob60Marketplace = 341.00;
fonteBob60Classico = 372.00;
fonteBob60Premium = 402.00;

fonteBob120Marketplace = 462.00;
fonteBob120Classico = 502.00;
fonteBob120Premium = 543.00;

fonteBob200Marketplace = 593.00;
fonteBob200Classico = 634.00;
fonteBob200Premium = 674.00;


fonteBaterryMeter50Marketplace = 412.00;
fonteBaterryMeter50Classico = 442.00;
fonteBaterryMeter50Premium = 472.00;

fonteBaterryMeter70Marketplace = 442.00;
fonteBaterryMeter70Classico = 473.00;
fonteBaterryMeter70Premium = 502.00;

fonteBaterryMeter100Marketplace = 523.00;
fonteBaterryMeter100Classico = 563.00;
fonteBaterryMeter100Premium = 603.00;

fonteBaterryMeter120Marketplace = 583.00;
fonteBaterryMeter120Classico = 623.00;
fonteBaterryMeter120Premium = 664.00;

fonteSmart50Marketplace = 442.00;
fonteSmart50Classico = 473.00;
fonteSmart50Premium = 502.00;

fonteSmart70Marketplace = 473.00;
fonteSmart70Classico = 502.00;
fonteSmart70Premium = 532.00;

fonteSmart100Marketplace = 553.00;
fonteSmart100Classico = 593.00;
fonteSmart100Premium = 633.00;

fonteSmart120Marketplace = 614.00;
fonteSmart120Classico = 654.00;
fonteSmart120Premium = 693.00;

fonteSmart160Marketplace = 755.00;
fonteSmart160Classico = 794.00;
fonteSmart160Premium = 834.00;

fonteSmart200MonoMarketplace = 745.00;
fonteSmart200MonoClassico = 784.00;
fonteSmart200MonoPremium = 825.00;

fonteSmart200Marketplace = 775.00;
fonteSmart200Classico = 814.00;
fonteSmart200Premium = 854.00;

fonteHeavyDuty220Marketplace = 865.00;
fonteHeavyDuty220Classico = 905.00;
fonteHeavyDuty220Premium = 946.00;

fonte30Marketplace = 582.00;
fonte30Classico = 624.00;
fonte30Premium = 664.00;

fonte70Marketplace = 753.00;
fonte70Classico = 805.00;
fonte70Premium = 854.00;

fonte100Marketplace = 924.00;
fonte100Classico = 986.00;
fonte100Premium = 1046.00;

ConversorDeTensao30AMarketplace = 411.00;
ConversorDeTensao30AClassico = 452.00;
ConversorDeTensao30APremium = 492.00;

ConversorDeTensao60AMarketplace = 764.00;
ConversorDeTensao60AClassico = 805.00;
ConversorDeTensao60APremium = 854.00;

ConversorDeTensao120AMarketplace = 995.00;
ConversorDeTensao120AClassico = 1036.00;
ConversorDeTensao120APremium = 1127.00;

ConversorDeTensao240AMarketplace = 1711.00;
ConversorDeTensao240AClassico = 1761.00;
ConversorDeTensao240APremium = 1912.00;

CarregadorDeBateriasCharger60AMarketplace = 643.00;
CarregadorDeBateriasCharger60AClassico = 673.00;
CarregadorDeBateriasCharger60APremium = 734.00;


# if os.path.exists("dados_scrapy.docx"):
#     doc = Document("dados_scrapy.docx")
# else:
doc = Document()

def extract_price(response):
  price_selectors = [
      '//*[@id="price"]/div/div[1]/div[1]/span[1]/span/span[2]/text()',
      '//html/body/main/div[2]/div[5]/div/div[1]/div/div[1]/div/div[@class="ui-pdp-container__row ui-pdp-container__row--price"]/div/div[1]/div[1]/span/span/span[2]/text()',
      '//*[@id="ui-pdp-main-container"]/div[1]/div/div[1]/div[2]/div[3]/div[1]/div[1]/span/span/span[2]/text()',
      '//*[@id="ui-pdp-main-container"]/div[1]/div/div[1]/div[2]/div[2]/div[1]/div[1]/span[1]/span/span[2]/text()'
  ]
  
  for selector in price_selectors:
    price = response.xpath(selector).get()
    if price:
      price = price.replace('.', '')
      decimal_selector = selector.replace("span[2]/text()", "") + 'span[@class="andes-money-amount__cents andes-money-amount__cents--superscript-36"]/text()'
      price_decimal = response.xpath(decimal_selector).get()
      
      if price_decimal:
        return float(f"{price}.{price_decimal}")
      else:
        try:
          return float(price)
        except ValueError:
          pass

  return None  


class MlSpider(scrapy.Spider):
    option_selected = ""
    option_selected_new = ""
    name = 'ml'
    start_urls = ["https://lista.mercadolivre.com.br/fonte-jfa"]
    
    def __init__(self, palavra=None, cookie=None, *args, **kwargs):
        super(MlSpider, self).__init__(*args, **kwargs)
        self.palavra = palavra
        self.cookie = cookie
    
    
    def parse(self, response, **kwargs):
        yield from self.parse_all(self.palavra)
        
        
    def parse_catalog(self, search_coming, headers):
        request = requests.get(
            f"https://app.nubimetrics.com/api/search/catalogs?order=relevance&search_filters=condition%3Dnew@&seller_id=1242763049&site_id=MLB&to_search={search_coming}&typeFind=q",
            headers=headers
        )
        request.raise_for_status()
        data = request.json()
        data = data.get('data', [])
        for i in data:
            name = i.get("Name")
            loja = i.get('BuyBoxWinner', {}).get("Nickname")
            listing_type = i.get('BuyBoxWinner', {}).get("ListingTypeId")
            if name and loja and listing_type:
                new_name = name.lower();
                new_name = unidecode.unidecode(new_name)
                if self.option_selected == "FONTE 40A":       
                    if "bob" not in new_name and "lite" not in new_name and "light" not in new_name  and "controle" not in new_name and 'jfa' in new_name:
                        if "40a" in new_name or "40" in new_name or "40 amperes" in new_name or "40amperes" in new_name or "36a" in new_name or "36" in new_name or "36 amperes" in new_name or "36amperes" in new_name:
                            yield scrapy.Request(url=i.get('Permalink'), callback=self.parse_product, meta={'name': name, 'loja': loja, 'listing_type': listing_type})          
                            for item in i.get('Sellers', []):
                                loja = item.get("Nickname")
                                name = i.get("Name")
                                listing_type = item.get("ListingTypeId")
                                if loja and name and listing_type:
                                    yield scrapy.Request(url=f'{i.get('Permalink')}?pdp_filters=item_id:{item.get('ItemId')}', callback=self.parse_product, meta={'name': name, 'loja': loja, 'listing_type': listing_type})
                                else:
                                    print("Missing loja or name for item:", i.get("Name"))
    
    def parse_all(self, option_function):
        self.option_selected = option_function
        self.option_selected_new = option_function
        search = ""
        if self.option_selected == "Fonte Usina Bob 60A":
            search = "Fonte Usina Bob 60A"
        if self.option_selected == "Fonte Usina Bob 120A":
            search = "Fonte Usina Bob 120A"
        elif self.option_selected == "Fonte Usina Bob 200A":
            search = "Fonte Usina Bob 200A"
        elif self.option_selected == "Fonte Usina Battery Meter 50A":
            search = "Fonte Usina Battery Meter 50A"
        elif self.option_selected == "Fonte Usina Battery Meter 70A":
            search = "Fonte Usina Battery Meter 70A"
        elif self.option_selected == "Fonte Usina Battery Meter 100A":
            search = "Fonte Usina Battery Meter 100A"
        elif self.option_selected == "Fonte Usina Battery Meter 120A":
            search = "Fonte Usina Battery Meter 120A"
        elif self.option_selected == "Fonte Usina Smart 50A":
            search = "Fonte Usina Smart 50A"
        elif self.option_selected == "Fonte Usina Smart 70A":
            search = "Fonte Usina Smart 70A"
        elif self.option_selected == "Fonte Usina Smart 100A":
            search = "Fonte Usina Smart 100A"
        elif self.option_selected == "Fonte Usina Smart 120A":
            search = "Fonte Usina Smart 120A"
        elif self.option_selected == "Fonte Usina Smart 160A":
            search = "Fonte Usina Smart 160A"
        elif self.option_selected == "Fonte Usina Smart 200A MONO":
            search = "Fonte Usina Smart 200A mono"
        elif self.option_selected == "Fonte Usina Smart 200A":
            search = "Fonte Usina Smart 200A"
        elif self.option_selected == "Fonte Usina 220A":
            search = "Fonte Usina 220A"
        elif self.option_selected == "Fonte Usina 30A":
            search = "Fonte Usina 30A"
        elif self.option_selected == "Fonte Usina 70A":
            search = "Fonte Usina 70A"
        elif self.option_selected == "Fonte Usina 100A":
            search = "Fonte Usina 100A"
        elif self.option_selected == "Conversor de Tensao 30A":
            search = "Conversor+de+Tensao+Usina+30A"
        elif self.option_selected == "Conversor de Tensao 60A":
            search = "Conversor+de+Tensao+Usina+60A"
        elif self.option_selected == "Conversor de Tensao 120A":
            search = "Conversor+de+Tensao+Usina+120A"
        elif self.option_selected == "Conversor de Tensao 240A":
            search = "Conversor+de+Tensao+Usina+240A"
        elif self.option_selected == "Carregador de Baterias Charger 60A":
            search = "Carregador+de+Baterias+Charger+Usina+60A"


        search = search.lower()
        search = search.replace(" ", "%20")
        
        search_catalog = ""
        if self.option_selected == "FONTE 200 MONO":
            search_catalog = "fonte 200a mono"
        search_catalog = search_catalog.replace(" ", "%20")
        
        headers = {
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.7",
            "Accept-Encoding": "gzip, deflate, br",
            "Accept-Language": "pt-BR,pt;q=0.9,en-US;q=0.8,en;q=0.7",
            "Cache-Control": "max-age=0",
            "Connection": "keep-alive",
            "Cookie": self.cookie,
            "Host": "app.nubimetrics.com",
            "Sec-Fetch-Dest": "document",
            "Sec-Fetch-Mode": "navigate",
            "Sec-Fetch-Site": "none",
            "Sec-Fetch-User": "?1",
            "Upgrade-Insecure-Requests": "1",
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36",
            "sec-ch-ua": "\"Google Chrome\";v=\"125\", \"Chromium\";v=\"125\", \"Not.A/Brand\";v=\"24\"",
            "sec-ch-ua-mobile": "?0",
            "sec-ch-ua-platform": "\"Windows\""
        }
        
        # yield from self.parse_catalog(search_catalog, headers)
                
        
        request = requests.get(
            f"https://app.nubimetrics.com/api/search/items?attributes=results,available_filters,paging,filters&buying_mode=buy_it_now&exportData=false&isControlPrice=false&language=pt_BR&limit=200&offset=0&order=price_asc&pvp=0&search_filters=condition%3Dnew@&seller_id=1242763049&site_id=MLB&to_search={search}&typeSearch=q",
            headers=headers
        )
        
        request.raise_for_status()
        data = request.json()
        paging_data = data.get('data', {}).get('paging', {}) 

        total = paging_data.get('total') 
        offset = paging_data.get('offset') 
        limit = paging_data.get('limit') 
        while offset < total:
            request = requests.get(
                f"https://app.nubimetrics.com/api/search/items?attributes=results,available_filters,paging,filters&buying_mode=buy_it_now&exportData=false&isControlPrice=false&language=pt_BR&limit=200&offset={offset}&order=price_asc&pvp=0&search_filters=condition%3Dnew@&seller_id=1242763049&site_id=MLB&to_search={search}&typeSearch=q",
                headers=headers
            )
            request.raise_for_status()
            data = request.json()
            paging_data = data.get('data', {}).get('paging', {}) 
            total = paging_data.get('total') 
            offset = paging_data.get('offset') 
            limit = paging_data.get('limit') 
            offset = offset + limit;
            
            for item in data.get('data', {}).get('results', []):
                name = item.get('title')
                id = item.get('id')
                permalink = item.get("permalink");
                if "/p/" in permalink:
                    url = permalink + f"?pdp_filters=item_id:{id}"
                else:
                    url = permalink
                listing_type = item.get('listing_type_id')
                loja = item.get('sellernickname')
                price = item.get('price')
                new_name = name.lower();
                new_name = unidecode.unidecode(new_name)
                if self.option_selected == "Fonte Usina Bob 60A":       
                    print(new_name)
                    if "bob" in new_name and "usina" in new_name and "smart" not in new_name and "samrt" not in new_name and "battery" not in new_name and "meter" not in new_name and "24v" not in new_name:          
                        if "60a" in new_name or "60" in new_name or "60 amperes" in new_name or "60amperes" in new_name:
                            yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type})
                            yield scrapy.Request(url='https://www.radicalsom.com.br/fonte-usina-60a-bob_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.bestonline.com.br/fonte-usina-60a-bob_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.shoppratico.com.br/fonte-usina-60a-bob_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.renovonline.com.br/fonte-usina-60a-bob_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.lsdistribuidora.com.br/fonte-usina-60a-bob_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})

                if self.option_selected == "Fonte Usina Bob 120A":       
                    if "bob" in new_name and "usina" in new_name and "smart" not in new_name and "samrt" not in new_name and "battery" not in new_name and "meter" not in new_name and "24v" not in new_name:
                        if "120a" in new_name or "120" in new_name or "120 amperes" in new_name or "120amperes" in new_name:
                            yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type})
                            yield scrapy.Request(url='https://www.radicalsom.com.br/fonte-usina-120a-bob_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.bestonline.com.br/fonte-usina-120a-bob_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.shoppratico.com.br/fonte-usina-120a-bob_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.renovonline.com.br/fonte-usina-120a-bob_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.lsdistribuidora.com.br/fonte-usina-120a-bob_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})
                            
                if self.option_selected == "Fonte Usina Bob 200A":       
                    if "bob" in new_name and "usina" in new_name and "smart" not in new_name and "samrt" not in new_name and "battery" not in new_name and "meter" not in new_name and "24v" not in new_name and "24v" not in new_name:
                        if "200a" in new_name or "200" in new_name or "200 amperes" in new_name or "200amperes" in new_name:
                            yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type})
                            yield scrapy.Request(url='https://www.radicalsom.com.br/fonte-usina-200a-bob_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.bestonline.com.br/fonte-usina-200a-bob_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.shoppratico.com.br/fonte-usina-200a-bob_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.renovonline.com.br/fonte-usina-200a-bob_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.lsdistribuidora.com.br/fonte-usina-200a-bob_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})
                            
                if self.option_selected == "Fonte Usina Battery Meter 50A":       
                    if "usina" in new_name and "battery" in new_name and "meter" in new_name and "bob" not in new_name and "24v" not in new_name:
                        if "50a" in new_name or "50" in new_name or "50 amperes" in new_name or "50amperes" in new_name:
                            yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type})
                            yield scrapy.Request(url='https://www.radicalsom.com.br/fonte-usina-50a-battery-meter_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.bestonline.com.br/fonte-usina-50a-battery-meter_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.shoppratico.com.br/fonte-usina-50a-battery-meter_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.renovonline.com.br/fonte-usina-50a-battery-meter_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.lsdistribuidora.com.br/fonte-usina-50a-battery-meter_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})
                            
                if self.option_selected == "Fonte Usina Battery Meter 70A":       
                    if "usina" in new_name and "battery" in new_name and "meter" in new_name and "bob" not in new_name and "24v" not in new_name:
                        if "50a" in new_name or "50" in new_name or "50 amperes" in new_name or "50amperes" in new_name:
                            yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type})
                            yield scrapy.Request(url='https://www.radicalsom.com.br/fonte-usina-50a-battery-meter_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.bestonline.com.br/fonte-usina-50a-battery-meter_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.shoppratico.com.br/fonte-usina-50a-battery-meter_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.renovonline.com.br/fonte-usina-50a-battery-meter_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.lsdistribuidora.com.br/fonte-usina-50a-battery-meter_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})
                            
                if self.option_selected == "Fonte Usina Battery Meter 100A":       
                    if "usina" in new_name and "battery" in new_name and "meter" in new_name and "bob" not in new_name and "24v" not in new_name:
                        if "100a" in new_name or "100" in new_name or "100 amperes" in new_name or "100amperes" in new_name:
                            yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type})
                            yield scrapy.Request(url='https://www.radicalsom.com.br/fonte-usina-100a-battery-meter_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.bestonline.com.br/fonte-usina-100a-battery-meter_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.shoppratico.com.br/fonte-usina-100a-battery-meter_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.renovonline.com.br/fonte-usina-100a-battery-meter_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.lsdistribuidora.com.br/fonte-usina-100a-battery-meter_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})
                            
                if self.option_selected == "Fonte Usina Battery Meter 120A":       
                    if "usina" in new_name and "battery" in new_name and "meter" in new_name and "bob" not in new_name and "24v" not in new_name:
                        if "120a" in new_name or "120" in new_name or "120 amperes" in new_name or "120amperes" in new_name:
                            yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type})
                            yield scrapy.Request(url='https://www.radicalsom.com.br/fonte-usina-120a-battery-meter_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.bestonline.com.br/fonte-usina-120a-battery-meter_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.shoppratico.com.br/fonte-usina-120a-battery-meter_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.renovonline.com.br/fonte-usina-120a-battery-meter_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.lsdistribuidora.com.br/fonte-usina-120a-battery-meter_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})
                            
                if self.option_selected == "Fonte Usina Smart 50A":       
                    if "usina" in new_name and ("smart" in new_name or "samrt" in new_name) and "bob" not in new_name and "battery" not in new_name and "meter" not in new_name and "24v" not in new_name:
                        if "50a" in new_name or "50" in new_name or "50 amperes" in new_name or "50amperes" in new_name:
                            yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type})
                            yield scrapy.Request(url='https://www.radicalsom.com.br/fonte-usina-50a-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.bestonline.com.br/fonte-usina-50a-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.shoppratico.com.br/fonte-usina-50a-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.renovonline.com.br/fonte-usina-50a-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.lsdistribuidora.com.br/fonte-usina-50a-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})
                            
                if self.option_selected == "Fonte Usina Smart 70A":       
                    if "usina" in new_name and ("smart" in new_name or "samrt" in new_name) and "bob" not in new_name and "battery" not in new_name and "meter" not in new_name and "24v" not in new_name:
                        if "70a" in new_name or "70" in new_name or "70 amperes" in new_name or "70amperes" in new_name:
                            yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type})
                            yield scrapy.Request(url='https://www.radicalsom.com.br/fonte-usina-70a-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.bestonline.com.br/fonte-usina-70a-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.shoppratico.com.br/fonte-usina-70a-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.renovonline.com.br/fonte-usina-70a-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.lsdistribuidora.com.br/fonte-usina-70a-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})
                            
                if self.option_selected == "Fonte Usina Smart 100A":       
                    if "usina" in new_name and ("smart" in new_name or "samrt" in new_name) and "bob" not in new_name and "battery" not in new_name and "meter" not in new_name and "24v" not in new_name:
                        if "100a" in new_name or "100" in new_name or "100 amperes" in new_name or "100amperes" in new_name:
                            yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type})
                            yield scrapy.Request(url='https://www.radicalsom.com.br/fonte-usina-100a-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.bestonline.com.br/fonte-usina-100a-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.shoppratico.com.br/fonte-usina-100a-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.renovonline.com.br/fonte-usina-100a-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.lsdistribuidora.com.br/fonte-usina-100a-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})
                            
                if self.option_selected == "Fonte Usina Smart 120A":       
                    if "usina" in new_name and ("smart" in new_name or "samrt" in new_name) and "bob" not in new_name and "battery" not in new_name and "meter" not in new_name and "24v" not in new_name:
                        if "120a" in new_name or "120" in new_name or "120 amperes" in new_name or "120amperes" in new_name:
                            yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type})
                            yield scrapy.Request(url='https://www.radicalsom.com.br/fonte-usina-120a-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.bestonline.com.br/fonte-usina-120a-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.shoppratico.com.br/fonte-usina-120a-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.renovonline.com.br/fonte-usina-120a-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.lsdistribuidora.com.br/fonte-usina-120a-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})
                            
                if self.option_selected == "Fonte Usina Smart 160A":       
                    if "usina" in new_name and ("smart" in new_name or "samrt" in new_name) and "bob" not in new_name and "battery" not in new_name and "meter" not in new_name and "24v" not in new_name:
                        if "160a" in new_name or "160" in new_name or "160 amperes" in new_name or "160amperes" in new_name:
                            yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type})
                            yield scrapy.Request(url='https://www.radicalsom.com.br/fonte-usina-160a-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.bestonline.com.br/fonte-usina-160a-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.shoppratico.com.br/fonte-usina-160a-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.renovonline.com.br/fonte-usina-160a-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.lsdistribuidora.com.br/fonte-usina-160a-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})
                            
                if self.option_selected == "Fonte Usina Smart 200A MONO":       
                    if "usina" in new_name and ("smart" in new_name or "samrt" in new_name) and ("mono" in new_name or "220v" in new_name or "monovolt" in new_name) and "bob" not in new_name and "battery" not in new_name and "meter" not in new_name and "24v" not in new_name:
                        if "200a" in new_name or "200" in new_name or "200 amperes" in new_name or "200amperes" in new_name:
                            yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type})
                            yield scrapy.Request(url='https://www.radicalsom.com.br/fonte-usina-200a-mono-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.bestonline.com.br/fonte-usina-200a-mono-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.shoppratico.com.br/fonte-usina-200a-mono-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.renovonline.com.br/fonte-usina-200a-mono-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.lsdistribuidora.com.br/fonte-usina-200a-mono-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})
                            
                if self.option_selected == "Fonte Usina Smart 200A":       
                    if "usina" in new_name and ("samrt" in new_name or "samrt" in new_name) and "bob" not in new_name and "battery" not in new_name and "meter" not in new_name and "24v" not in new_name:
                        if "200a" in new_name or "200" in new_name or "200 amperes" in new_name or "200amperes" in new_name:
                            yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type})
                            yield scrapy.Request(url='https://www.radicalsom.com.br/fonte-usina-200a-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.bestonline.com.br/fonte-usina-200a-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.shoppratico.com.br/fonte-usina-200a-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.renovonline.com.br/fonte-usina-200a-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.lsdistribuidora.com.br/fonte-usina-200a-smart_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})
                            
                if self.option_selected == "Fonte Usina 220A":       
                    if "usina" in new_name and "bob" not in new_name and "battery" not in new_name and "meter" not in new_name and "smart" not in new_name and "samrt" not in new_name and "24v" not in new_name:
                        if "220a" in new_name or "220" in new_name or "220 amperes" in new_name or "220amperes" in new_name:
                            yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type})
                            yield scrapy.Request(url='https://www.radicalsom.com.br/fonte-usina-220a_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.bestonline.com.br/fonte-usina-220a_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.shoppratico.com.br/fonte-usina-220a_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.renovonline.com.br/fonte-usina-220a_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.lsdistribuidora.com.br/fonte-usina-220a_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})
                            
                if self.option_selected == "Fonte Usina 30A":       
                    if "usina" in new_name and "bob" not in new_name and "battery" not in new_name and "meter" not in new_name and "smart" not in new_name and "samrt" not in new_name and "24v" in new_name:
                        if "30a" in new_name or "30" in new_name or "30 amperes" in new_name or "30amperes" in new_name:
                            yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type})
                            yield scrapy.Request(url='https://www.radicalsom.com.br/fonte-usina-30a_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.bestonline.com.br/fonte-usina-30a_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.shoppratico.com.br/fonte-usina-30a_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.renovonline.com.br/fonte-usina-30a_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.lsdistribuidora.com.br/fonte-usina-30a_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})
                            
                if self.option_selected == "Fonte Usina 70A":       
                    if "usina" in new_name and "bob" not in new_name and "battery" not in new_name and "meter" not in new_name and "smart" not in new_name and "samrt" not in new_name and "24v" in new_name:
                        if "70a" in new_name or "70" in new_name or "70 amperes" in new_name or "70amperes" in new_name:
                            yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type})
                            yield scrapy.Request(url='https://www.radicalsom.com.br/fonte-usina-70a_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.bestonline.com.br/fonte-usina-70a_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.shoppratico.com.br/fonte-usina-70a_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.renovonline.com.br/fonte-usina-70a_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.lsdistribuidora.com.br/fonte-usina-70a_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})
                            
                if self.option_selected == "Fonte Usina 100A":       
                    if "usina" in new_name and "bob" not in new_name and "battery" not in new_name and "meter" not in new_name and "smart" not in new_name and "samrt" not in new_name and "24v" in new_name:
                        if "100a" in new_name or "100" in new_name or "100 amperes" in new_name or "100amperes" in new_name:
                            yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type})
                            yield scrapy.Request(url='https://www.radicalsom.com.br/fonte-usina-100a_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.bestonline.com.br/fonte-usina-100a_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.shoppratico.com.br/fonte-usina-100a_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.renovonline.com.br/fonte-usina-100a_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.lsdistribuidora.com.br/fonte-usina-100a_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})
                            
                if self.option_selected == "Conversor de Tensao 30A":       
                    if "usina" in new_name and "conversor" in new_name:
                        if "30a" in new_name or "30" in new_name or "30 amperes" in new_name or "30amperes" in new_name:
                            yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type})
                            yield scrapy.Request(url='https://www.radicalsom.com.br/conversor_de_tensao_30a_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.bestonline.com.br/conversor_de_tensao_30a_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.shoppratico.com.br/conversor_de_tensao_30a_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.renovonline.com.br/conversor_de_tensao_30a_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.lsdistribuidora.com.br/conversor_de_tensao_30a_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})
                            
                if self.option_selected == "Conversor de Tensao 60A":       
                    if "usina" in new_name and "conversor" in new_name:
                        if "60a" in new_name or "60" in new_name or "60 amperes" in new_name or "60amperes" in new_name:
                            yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type})
                            yield scrapy.Request(url='https://www.radicalsom.com.br/conversor_de_tensao_60a_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.bestonline.com.br/conversor_de_tensao_60a_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.shoppratico.com.br/conversor_de_tensao_60a_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.renovonline.com.br/conversor_de_tensao_60a_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.lsdistribuidora.com.br/conversor_de_tensao_60a_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})
                                                        
                if self.option_selected == "Conversor de Tensao 120A":       
                    if "usina" in new_name and "conversor" in new_name:
                        if "120a" in new_name or "120" in new_name or "120 amperes" in new_name or "120amperes" in new_name:
                            yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type})
                            yield scrapy.Request(url='https://www.radicalsom.com.br/conversor_de_tensao_120a_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.bestonline.com.br/conversor_de_tensao_120a_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.shoppratico.com.br/conversor_de_tensao_120a_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.renovonline.com.br/conversor_de_tensao_120a_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.lsdistribuidora.com.br/conversor_de_tensao_120a_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})

                if self.option_selected == "Conversor de Tensao 240A":       
                    if "usina" in new_name and "conversor" in new_name:
                        if "240a" in new_name or "240" in new_name or "240 amperes" in new_name or "240amperes" in new_name:
                            yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type})
                            yield scrapy.Request(url='https://www.radicalsom.com.br/conversor_de_tensao_240a_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.bestonline.com.br/conversor_de_tensao_240a_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.shoppratico.com.br/conversor_de_tensao_240a_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.renovonline.com.br/conversor_de_tensao_240a_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.lsdistribuidora.com.br/conversor_de_tensao_240a_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})
                                                    
                if self.option_selected == "Carregador de Baterias Charger 60A":       
                    if "usina" in new_name and "charger" in new_name and "24v" not in new_name:
                        if "60a" in new_name or "60" in new_name or "60 amperes" in new_name or "60amperes" in new_name:
                            yield scrapy.Request(url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type})
                            yield scrapy.Request(url='https://www.radicalsom.com.br/carregador_baterias_charger_60a_OrderId_PRICE_NoIndex_True', callback=self.parse_radicalson, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.bestonline.com.br/carregador_baterias_charger_60a_OrderId_PRICE_NoIndex_True', callback=self.parse_bestonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.shoppratico.com.br/carregador_baterias_charger_60a_OrderId_PRICE_NoIndex_True', callback=self.parse_shoppratico, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.renovonline.com.br/carregador_baterias_charger_60a_OrderId_PRICE_NoIndex_True', callback=self.parse_renovonline, meta={'name': name, 'loja': loja, 'price':price})
                            yield scrapy.Request(url='https://www.lsdistribuidora.com.br/carregador_baterias_charger_60a_OrderId_PRICE_NoIndex_True', callback=self.parse_lsdistribuidora, meta={'name': name, 'loja': loja, 'price':price})
                                                    
                
    def parse_product(self, response):
        name = response.meta['name']
        loja = response.meta['loja']
        listing_type = response.meta['listing_type']

        self.option_selected_new = self.option_selected
        price = extract_price(response)
            
        if not price:
            return        
        new_price_float = price   
        
            
        parcelado = response.xpath('//*[@id="pricing_price_subtitle"]/text()').get()
        if parcelado:
            parcelado = int(parcelado.replace("x", "").strip())
        else:
            parcelado = 0
        
        other_price = response.xpath('//*[@id="pricing_price_subtitle"]/span[2]/span/span[2]/text()').get()
        other_price_decimal = response.xpath('//*[@id="pricing_price_subtitle"]/span[2]/span/span[4]/text()').get()
        
        if other_price and other_price_decimal:
            new_price_other = f"{other_price},{other_price_decimal}"
        elif other_price:
            new_price_other = other_price
        else:
            new_price_other = '0'

        try:
            new_price_other_float = float(new_price_other.replace('.', '').replace(',', '.'))
            new_price_other_float = round((new_price_other_float * parcelado), 3)
        except ValueError:
            new_price_other_float = 0.0
            

        target_id = response.xpath('//*[@id="denounce"]/div/p/span/text()').get()
        if target_id:
            target_id = target_id.replace("#", "")
    
        if listing_type:
            if listing_type == "gold_pro":
                tipo = "Premium"
            else:
                tipo = "Clássico" 
        else:
            tipo = "Clássico" 
        
            
        if tipo == "Clássico" and new_price_float:
            if self.option_selected_new == "Fonte Usina Bob 60A" and new_price_float >= fonteBob60Classico:
                return;
            elif self.option_selected_new == "Fonte Usina Bob 120A" and new_price_float >= fonteBob120Classico:
                return;
            elif self.option_selected_new == "Fonte Usina Bob 200A" and new_price_float >= fonteBob200Classico:
                return;
            elif self.option_selected_new == "Fonte Usina Battery Meter 50A" and new_price_float >= fonteBaterryMeter50Classico:
                return;
            elif self.option_selected_new == "Fonte Usina Battery Meter 70A" and new_price_float >= fonteBaterryMeter70Classico:
                return;
            elif self.option_selected_new == "Fonte Usina Battery Meter 100A" and new_price_float >= fonteBaterryMeter100Classico:
                return;
            elif self.option_selected_new == "Fonte Usina Battery Meter 120A" and new_price_float >= fonteBaterryMeter120Classico:
                return;
            elif self.option_selected_new == "Fonte Usina Smart 50A" and new_price_float >= fonteSmart50Classico:
                return;
            elif self.option_selected_new == "Fonte Usina Smart 70A" and new_price_float >= fonteSmart70Classico:
                return;
            elif self.option_selected_new == "Fonte Usina Smart 100A" and new_price_float >= fonteSmart100Classico:
                return;
            elif self.option_selected_new == "Fonte Usina Smart 120A" and new_price_float >= fonteSmart120Classico:
                return;
            elif self.option_selected_new == "Fonte Usina Smart 160A" and new_price_float >= fonteSmart160Classico:
                return;
            elif self.option_selected_new == "Fonte Usina Smart 200A MONO" and new_price_float >= fonteSmart200MonoClassico:
                return;
            elif self.option_selected_new == "Fonte Usina Smart 200A" and new_price_float >= fonteSmart200Classico:
                return;
            elif self.option_selected_new == "Fonte Usina 220A" and new_price_float >= fonteHeavyDuty220Classico:
                return;
            elif self.option_selected_new == "Fonte Usina 30A" and new_price_float >= fonte30Classico:
                return;
            elif self.option_selected_new == "Fonte Usina 70A" and new_price_float >= fonte70Classico:
                return;
            elif self.option_selected_new == "Fonte Usina 100A" and new_price_float >= fonte100Classico:
                return;
            elif self.option_selected_new == "Conversor de Tensao 30A" and new_price_float >= ConversorDeTensao30AClassico:
                return;
            elif self.option_selected_new == "Conversor de Tensao 60A" and new_price_float >= ConversorDeTensao60AClassico:
                return;
            elif self.option_selected_new == "Conversor de Tensao 120A" and new_price_float >= ConversorDeTensao120AClassico:
                return;
            elif self.option_selected_new == "Conversor de Tensao 240A" and new_price_float >= ConversorDeTensao240AClassico:
                return;
            elif self.option_selected_new == "Carregador de Baterias Charger 60A" and new_price_float >= CarregadorDeBateriasCharger60AClassico:
                return;
        elif tipo == "Premium" and new_price_float:
            if self.option_selected_new == "Fonte Usina Bob 60A" and new_price_float >= fonteBob60Premium:
                return;
            elif self.option_selected_new == "Fonte Usina Bob 120A" and new_price_float >= fonteBob120Premium:
                return;
            elif self.option_selected_new == "Fonte Usina Bob 200A" and new_price_float >= fonteBob200Premium:
                return;
            elif self.option_selected_new == "Fonte Usina Battery Meter 50A" and new_price_float >= fonteBaterryMeter50Premium:
                return;
            elif self.option_selected_new == "Fonte Usina Battery Meter 70A" and new_price_float >= fonteBaterryMeter70Premium:
                return;
            elif self.option_selected_new == "Fonte Usina Battery Meter 100A" and new_price_float >= fonteBaterryMeter100Premium:
                return;
            elif self.option_selected_new == "Fonte Usina Battery Meter 120A" and new_price_float >= fonteBaterryMeter120Premium:
                return;
            elif self.option_selected_new == "Fonte Usina Smart 50A" and new_price_float >= fonteSmart50Premium:
                return;
            elif self.option_selected_new == "Fonte Usina Smart 70A" and new_price_float >= fonteSmart70Premium:
                return;
            elif self.option_selected_new == "Fonte Usina Smart 100A" and new_price_float >= fonteSmart100Premium:
                return;
            elif self.option_selected_new == "Fonte Usina Smart 120A" and new_price_float >= fonteSmart120Premium:
                return;
            elif self.option_selected_new == "Fonte Usina Smart 160A" and new_price_float >= fonteSmart160Premium:
                return;
            elif self.option_selected_new == "Fonte Usina Smart 200A MONO" and new_price_float >= fonteSmart200MonoPremium:
                return;
            elif self.option_selected_new == "Fonte Usina Smart 200A" and new_price_float >= fonteSmart200Premium:
                return;
            elif self.option_selected_new == "Fonte Usina 220A" and new_price_float >= fonteHeavyDuty220Premium:
                return;
            elif self.option_selected_new == "Fonte Usina 30A" and new_price_float >= fonte30Premium:
                return;
            elif self.option_selected_new == "Fonte Usina 70A" and new_price_float >= fonte70Premium:
                return;
            elif self.option_selected_new == "Fonte Usina 100A" and new_price_float >= fonte100Premium:
                return;
            elif self.option_selected_new == "Conversor de Tensao 30A" and new_price_float >= ConversorDeTensao30APremium:
                return;
            elif self.option_selected_new == "Conversor de Tensao 60A" and new_price_float >= ConversorDeTensao60APremium:
                return;
            elif self.option_selected_new == "Conversor de Tensao 120A" and new_price_float >= ConversorDeTensao120APremium:
                return;
            elif self.option_selected_new == "Conversor de Tensao 240A" and new_price_float >= ConversorDeTensao240APremium:
                return;
            elif self.option_selected_new == "Carregador de Baterias Charger 60A" and new_price_float >= CarregadorDeBateriasCharger60APremium:
                return;


        location_url = f'https://www.mercadolivre.com.br/perfil/{loja.replace(" ", "+")}'

        yield scrapy.Request(url=location_url, callback=self.parse_location, meta={'url': response.url, 'name': name, 'price': new_price_float, 'qtde_parcelado': parcelado, 'price_parcelado': new_price_other_float, 'loja': loja, 'tipo': tipo })


    def finish(self, total_price, url, nomeFonte, loja, lugar):
        if self.option_selected_new == "Fonte Usina Bob 60A" and total_price >= fonteBob60Marketplace:
            return;
        elif self.option_selected_new == "Fonte Usina Bob 120A" and total_price >= fonteBob120Marketplace:
            return;
        elif self.option_selected_new == "Fonte Usina Bob 200A" and total_price >= fonteBob200Marketplace:
            return;
        elif self.option_selected_new == "Fonte Usina Battery Meter 50A" and total_price >= fonteBaterryMeter50Marketplace:
            return;
        elif self.option_selected_new == "Fonte Usina Battery Meter 70A" and total_price >= fonteBaterryMeter70Marketplace:
            return;
        elif self.option_selected_new == "Fonte Usina Battery Meter 100A" and total_price >= fonteBaterryMeter100Marketplace:
            return;
        elif self.option_selected_new == "Fonte Usina Battery Meter 120A" and total_price >= fonteBaterryMeter120Marketplace:
            return;
        elif self.option_selected_new == "Fonte Usina Smart 50A" and total_price >= fonteSmart50Marketplace:
            return;
        elif self.option_selected_new == "Fonte Usina Smart 70A" and total_price >= fonteSmart70Marketplace:
            return;
        elif self.option_selected_new == "Fonte Usina Smart 100A" and total_price >= fonteSmart100Marketplace:
            return;
        elif self.option_selected_new == "Fonte Usina Smart 120A" and total_price >= fonteSmart120Marketplace:
            return;
        elif self.option_selected_new == "Fonte Usina Smart 160A" and total_price >= fonteSmart160Marketplace:
            return;
        elif self.option_selected_new == "Fonte Usina Smart 200A MONO" and total_price >= fonteSmart200MonoMarketplace:
            return;
        elif self.option_selected_new == "Fonte Usina Smart 200A" and total_price >= fonteSmart200Marketplace:
            return;
        elif self.option_selected_new == "Fonte Usina 220A" and total_price >= fonteHeavyDuty220Marketplace:
            return;
        elif self.option_selected_new == "Fonte Usina 30A" and total_price >= fonte30Marketplace:
            return;
        elif self.option_selected_new == "Fonte Usina 70A" and total_price >= fonte70Marketplace:
            return;
        elif self.option_selected_new == "Fonte Usina 100A" and total_price >= fonte100Marketplace:
            return;
        elif self.option_selected_new == "Conversor de Tensao 30A" and total_price >= ConversorDeTensao30AMarketplace:
            return;
        elif self.option_selected_new == "Conversor de Tensao 60A" and total_price >= ConversorDeTensao60AMarketplace:
            return;
        elif self.option_selected_new == "Conversor de Tensao 120A" and total_price >= ConversorDeTensao120AMarketplace:
            return;
        elif self.option_selected_new == "Conversor de Tensao 240A" and total_price >= ConversorDeTensao240AMarketplace:
            return;
        elif self.option_selected_new == "Carregador de Baterias Charger 60A" and total_price >= CarregadorDeBateriasCharger60AMarketplace:
            return;
    
        parcelado = self.get_price_previsto("NA")

        doc.add_paragraph(f'Modelo: {self.option_selected_new}')
        doc.add_paragraph(f'URL: {url}')
        doc.add_paragraph(f'Nome: {nomeFonte}')
        doc.add_paragraph(f'Preço: {total_price}')
        doc.add_paragraph(f'Preço Previsto: {parcelado}')
        doc.add_paragraph(f'Loja: {loja}')
        doc.add_paragraph('Tipo: ')
        doc.add_paragraph(f'Lugar: {lugar}')
        doc.add_paragraph("--------------------------------------------------------------------")
        doc.add_paragraph('')
        doc.save(fr"dados/{self.option_selected_new}.docx") 
            
        yield {
            'url': url,
            'name': nomeFonte,
            'price': total_price,
            'loja': loja,
            'tipo': "",
            'lugar': lugar
        }


    def parse_radicalson(self, response):
        loja = "RADICALSOM"
        lugar = "Artur nogueira, São Paulo."
        for i in response.xpath('//*[@id="root-app"]/div/div[3]/section/ol/li'):
            nomeFonte = i.xpath('.//div/div/div/div/a/h2/text()').get()
            price = i.xpath('.//div/div/div[3]/div/div[1]/div/div/div/div/span/span[2]/text()').get()
            cents = i.xpath('.//div/div/div[3]/div/div[1]/div/div/div/div/span/span[4]/text()').get()
            url = i.xpath('.//div/div/div[3]/div[2]/a/@href').get()
            
            
            nomeFonte = nomeFonte.lower()
            nomeFonte = unidecode.unidecode(nomeFonte)
            if not cents:
                cents = 0
            if price:
                price = price.replace('.', '')
                total_price = float(f"{price}.{cents}")
            else:
                return
            if self.option_selected == "Fonte Usina Bob 60A":       
                if "bob" in nomeFonte and "usina" in nomeFonte and "smart" not in nomeFonte and "samrt" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "60a" in nomeFonte or "60" in nomeFonte or "60 amperes" in nomeFonte or "60amperes" in nomeFonte:
                            yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Bob 120A":       
                if "bob" in nomeFonte and "usina" in nomeFonte and "smart" not in nomeFonte and "samrt" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Bob 200A":       
                if "bob" in nomeFonte and "usina" in nomeFonte and "smart" not in nomeFonte and "samrt" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte and "12v" in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Battery Meter 50A":       
                if "usina" in nomeFonte and "battery" in nomeFonte and "meter" in nomeFonte and "bob" not in nomeFonte and "12v" in nomeFonte:
                    if "50a" in nomeFonte or "50" in nomeFonte or "50 amperes" in nomeFonte or "50amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Battery Meter 70A":       
                if "usina" in nomeFonte and "battery" in nomeFonte and "meter" in nomeFonte and "bob" not in nomeFonte and "12v" in nomeFonte:
                    if "50a" in nomeFonte or "50" in nomeFonte or "50 amperes" in nomeFonte or "50amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Battery Meter 100A":       
                if "usina" in nomeFonte and "battery" in nomeFonte and "meter" in nomeFonte and "bob" not in nomeFonte and "12v" in nomeFonte:
                    if "100a" in nomeFonte or "100" in nomeFonte or "100 amperes" in nomeFonte or "100amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Battery Meter 120A":       
                if "usina" in nomeFonte and "battery" in nomeFonte and "meter" in nomeFonte and "bob" not in nomeFonte and "12v" in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Smart 50A":       
                if "usina" in nomeFonte and ("smart" in nomeFonte or "samrt" in nomeFonte) and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "50a" in nomeFonte or "50" in nomeFonte or "50 amperes" in nomeFonte or "50amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Smart 70A":       
                if "usina" in nomeFonte and ("smart" in nomeFonte or "samrt" in nomeFonte) and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "70a" in nomeFonte or "70" in nomeFonte or "70 amperes" in nomeFonte or "70amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Smart 100A":       
                if "usina" in nomeFonte and ("smart" in nomeFonte or "samrt" in nomeFonte) and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "100a" in nomeFonte or "100" in nomeFonte or "100 amperes" in nomeFonte or "100amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Smart 120A":       
                if "usina" in nomeFonte and ("smart" in nomeFonte or "samrt" in nomeFonte) and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Smart 160A":       
                if "usina" in nomeFonte and ("smart" in nomeFonte or "samrt" in nomeFonte) and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "160a" in nomeFonte or "160" in nomeFonte or "160 amperes" in nomeFonte or "160amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Smart 200A MONO":       
                if "usina" in nomeFonte and ("smart" in nomeFonte or "samrt" in nomeFonte) and ("mono" in nomeFonte or "220v" in nomeFonte or "monovolt" in nomeFonte) and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Smart 200A":       
                if "usina" in nomeFonte and ("samrt" in nomeFonte or "samrt" in nomeFonte) and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina 220A":       
                if "usina" in nomeFonte and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "smart" not in nomeFonte and "samrt" not in nomeFonte and "12v" in nomeFonte:
                    if "220a" in nomeFonte or "220" in nomeFonte or "220 amperes" in nomeFonte or "220amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina 30A":       
                if "usina" in nomeFonte and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "smart" not in nomeFonte and "samrt" not in nomeFonte and "24v" in nomeFonte:
                    if "30a" in nomeFonte or "30" in nomeFonte or "30 amperes" in nomeFonte or "30amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina 70A":       
                if "usina" in nomeFonte and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "smart" not in nomeFonte and "samrt" not in nomeFonte and "24v" in nomeFonte:
                    if "70a" in nomeFonte or "70" in nomeFonte or "70 amperes" in nomeFonte or "70amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina 100A":       
                if "usina" in nomeFonte and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "smart" not in nomeFonte and "samrt" not in nomeFonte and "24v" in nomeFonte:
                    if "100a" in nomeFonte or "100" in nomeFonte or "100 amperes" in nomeFonte or "100amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Conversor de Tensao 30A":       
                if "usina" in nomeFonte and "conversor" in nomeFonte:
                    if "30a" in nomeFonte or "30" in nomeFonte or "30 amperes" in nomeFonte or "30amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Conversor de Tensao 60A":       
                if "usina" in nomeFonte and "conversor" in nomeFonte:
                    if "60a" in nomeFonte or "60" in nomeFonte or "60 amperes" in nomeFonte or "60amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Conversor de Tensao 120A":       
                if "usina" in nomeFonte and "conversor" in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Conversor de Tensao 240A":       
                if "usina" in nomeFonte and "conversor" in nomeFonte:
                    if "240a" in nomeFonte or "240" in nomeFonte or "240 amperes" in nomeFonte or "240amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Carregador de Baterias Charger 60A":       
                if "usina" in nomeFonte and "charger" in nomeFonte and "12v" in nomeFonte:
                    if "60a" in nomeFonte or "60" in nomeFonte or "60 amperes" in nomeFonte or "60amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                            
                        
                    
                    
    def parse_lsdistribuidora(self, response):
        loja = "LS DISTRIBUIDORA"
        lugar = "Elísio Medrado, Bahia"
        for i in response.xpath('//*[@id="root-app"]/div/div[3]/section/ol/li'):
            nomeFonte = i.xpath('.//div/div/div/div/a/h2/text()').get()
            price = i.xpath('.//div/div/div[3]/div/div[1]/div/div/div/div/span/span[2]/text()').get()
            cents = i.xpath('.//div/div/div[3]/div/div[1]/div/div/div/div/span/span[4]/text()').get()
            url = i.xpath('.//div/div/div[3]/div[2]/a/@href').get()
            nomeFonte = nomeFonte.lower()
            nomeFonte = unidecode.unidecode(nomeFonte)
            if not cents:
                cents = 0
            if price:
                price = price.replace('.', '')
                total_price = float(f"{price}.{cents}")
            if self.option_selected == "Fonte Usina Bob 60A":       
                if "bob" in nomeFonte and "usina" in nomeFonte and "smart" not in nomeFonte and "samrt" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "60a" in nomeFonte or "60" in nomeFonte or "60 amperes" in nomeFonte or "60amperes" in nomeFonte:
                            yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Bob 120A":       
                if "bob" in nomeFonte and "usina" in nomeFonte and "smart" not in nomeFonte and "samrt" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Bob 200A":       
                if "bob" in nomeFonte and "usina" in nomeFonte and "smart" not in nomeFonte and "samrt" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte and "12v" in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Battery Meter 50A":       
                if "usina" in nomeFonte and "battery" in nomeFonte and "meter" in nomeFonte and "bob" not in nomeFonte and "12v" in nomeFonte:
                    if "50a" in nomeFonte or "50" in nomeFonte or "50 amperes" in nomeFonte or "50amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Battery Meter 70A":       
                if "usina" in nomeFonte and "battery" in nomeFonte and "meter" in nomeFonte and "bob" not in nomeFonte and "12v" in nomeFonte:
                    if "50a" in nomeFonte or "50" in nomeFonte or "50 amperes" in nomeFonte or "50amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Battery Meter 100A":       
                if "usina" in nomeFonte and "battery" in nomeFonte and "meter" in nomeFonte and "bob" not in nomeFonte and "12v" in nomeFonte:
                    if "100a" in nomeFonte or "100" in nomeFonte or "100 amperes" in nomeFonte or "100amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Battery Meter 120A":       
                if "usina" in nomeFonte and "battery" in nomeFonte and "meter" in nomeFonte and "bob" not in nomeFonte and "12v" in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Smart 50A":       
                if "usina" in nomeFonte and ("smart" in nomeFonte or "samrt" in nomeFonte) and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "50a" in nomeFonte or "50" in nomeFonte or "50 amperes" in nomeFonte or "50amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Smart 70A":       
                if "usina" in nomeFonte and ("smart" in nomeFonte or "samrt" in nomeFonte) and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "70a" in nomeFonte or "70" in nomeFonte or "70 amperes" in nomeFonte or "70amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Smart 100A":       
                if "usina" in nomeFonte and ("smart" in nomeFonte or "samrt" in nomeFonte) and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "100a" in nomeFonte or "100" in nomeFonte or "100 amperes" in nomeFonte or "100amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Smart 120A":       
                if "usina" in nomeFonte and ("smart" in nomeFonte or "samrt" in nomeFonte) and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Smart 160A":       
                if "usina" in nomeFonte and ("smart" in nomeFonte or "samrt" in nomeFonte) and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "160a" in nomeFonte or "160" in nomeFonte or "160 amperes" in nomeFonte or "160amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Smart 200A MONO":       
                if "usina" in nomeFonte and ("smart" in nomeFonte or "samrt" in nomeFonte) and ("mono" in nomeFonte or "220v" in nomeFonte or "monovolt" in nomeFonte) and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Smart 200A":       
                if "usina" in nomeFonte and ("samrt" in nomeFonte or "samrt" in nomeFonte) and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina 220A":       
                if "usina" in nomeFonte and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "smart" not in nomeFonte and "samrt" not in nomeFonte and "12v" in nomeFonte:
                    if "220a" in nomeFonte or "220" in nomeFonte or "220 amperes" in nomeFonte or "220amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina 30A":       
                if "usina" in nomeFonte and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "smart" not in nomeFonte and "samrt" not in nomeFonte and "24v" in nomeFonte:
                    if "30a" in nomeFonte or "30" in nomeFonte or "30 amperes" in nomeFonte or "30amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina 70A":       
                if "usina" in nomeFonte and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "smart" not in nomeFonte and "samrt" not in nomeFonte and "24v" in nomeFonte:
                    if "70a" in nomeFonte or "70" in nomeFonte or "70 amperes" in nomeFonte or "70amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina 100A":       
                if "usina" in nomeFonte and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "smart" not in nomeFonte and "samrt" not in nomeFonte and "24v" in nomeFonte:
                    if "100a" in nomeFonte or "100" in nomeFonte or "100 amperes" in nomeFonte or "100amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Conversor de Tensao 30A":       
                if "usina" in nomeFonte and "conversor" in nomeFonte:
                    if "30a" in nomeFonte or "30" in nomeFonte or "30 amperes" in nomeFonte or "30amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Conversor de Tensao 60A":       
                if "usina" in nomeFonte and "conversor" in nomeFonte:
                    if "60a" in nomeFonte or "60" in nomeFonte or "60 amperes" in nomeFonte or "60amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Conversor de Tensao 120A":       
                if "usina" in nomeFonte and "conversor" in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Conversor de Tensao 240A":       
                if "usina" in nomeFonte and "conversor" in nomeFonte:
                    if "240a" in nomeFonte or "240" in nomeFonte or "240 amperes" in nomeFonte or "240amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Carregador de Baterias Charger 60A":       
                if "usina" in nomeFonte and "charger" in nomeFonte and "12v" in nomeFonte:
                    if "60a" in nomeFonte or "60" in nomeFonte or "60 amperes" in nomeFonte or "60amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

    
    def parse_bestonline(self, response):
        loja = "BESTONLINE"
        lugar = "Rosario, Santa Fe."
        for i in response.xpath('//li[@class="ui-search-layout__item shops__layout-item shops__layout-item ui-search-layout__stack"]'):
            nomeFonte = i.xpath('.//div/div/div/div/a/h2/text()').get()
            price = i.xpath('.//div/div/div[3]/div/div[1]/div/div/div/div/span/span[2]/text()').get()
            cents = i.xpath('.//div/div/div[3]/div/div[1]/div/div/div/div/span/span[4]/text()').get()
            url = i.xpath('.//div/div/div[3]/div[2]/a/@href').get()
            nomeFonte = nomeFonte.lower()
            nomeFonte = unidecode.unidecode(nomeFonte)
            if not cents:
                cents = 0
            if price:
                price = price.replace('.', '')
                total_price = float(f"{price}.{cents}")
            if self.option_selected == "Fonte Usina Bob 60A":       
                if "bob" in nomeFonte and "usina" in nomeFonte and "smart" not in nomeFonte and "samrt" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "60a" in nomeFonte or "60" in nomeFonte or "60 amperes" in nomeFonte or "60amperes" in nomeFonte:
                            yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Bob 120A":       
                if "bob" in nomeFonte and "usina" in nomeFonte and "smart" not in nomeFonte and "samrt" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Bob 200A":       
                if "bob" in nomeFonte and "usina" in nomeFonte and "smart" not in nomeFonte and "samrt" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte and "12v" in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Battery Meter 50A":       
                if "usina" in nomeFonte and "battery" in nomeFonte and "meter" in nomeFonte and "bob" not in nomeFonte and "12v" in nomeFonte:
                    if "50a" in nomeFonte or "50" in nomeFonte or "50 amperes" in nomeFonte or "50amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Battery Meter 70A":       
                if "usina" in nomeFonte and "battery" in nomeFonte and "meter" in nomeFonte and "bob" not in nomeFonte and "12v" in nomeFonte:
                    if "50a" in nomeFonte or "50" in nomeFonte or "50 amperes" in nomeFonte or "50amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Battery Meter 100A":       
                if "usina" in nomeFonte and "battery" in nomeFonte and "meter" in nomeFonte and "bob" not in nomeFonte and "12v" in nomeFonte:
                    if "100a" in nomeFonte or "100" in nomeFonte or "100 amperes" in nomeFonte or "100amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Battery Meter 120A":       
                if "usina" in nomeFonte and "battery" in nomeFonte and "meter" in nomeFonte and "bob" not in nomeFonte and "12v" in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Smart 50A":       
                if "usina" in nomeFonte and ("smart" in nomeFonte or "samrt" in nomeFonte) and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "50a" in nomeFonte or "50" in nomeFonte or "50 amperes" in nomeFonte or "50amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Smart 70A":       
                if "usina" in nomeFonte and ("smart" in nomeFonte or "samrt" in nomeFonte) and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "70a" in nomeFonte or "70" in nomeFonte or "70 amperes" in nomeFonte or "70amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Smart 100A":       
                if "usina" in nomeFonte and ("smart" in nomeFonte or "samrt" in nomeFonte) and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "100a" in nomeFonte or "100" in nomeFonte or "100 amperes" in nomeFonte or "100amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Smart 120A":       
                if "usina" in nomeFonte and ("smart" in nomeFonte or "samrt" in nomeFonte) and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Smart 160A":       
                if "usina" in nomeFonte and ("smart" in nomeFonte or "samrt" in nomeFonte) and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "160a" in nomeFonte or "160" in nomeFonte or "160 amperes" in nomeFonte or "160amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Smart 200A MONO":       
                if "usina" in nomeFonte and ("smart" in nomeFonte or "samrt" in nomeFonte) and ("mono" in nomeFonte or "220v" in nomeFonte or "monovolt" in nomeFonte) and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Smart 200A":       
                if "usina" in nomeFonte and ("samrt" in nomeFonte or "samrt" in nomeFonte) and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina 220A":       
                if "usina" in nomeFonte and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "smart" not in nomeFonte and "samrt" not in nomeFonte and "12v" in nomeFonte:
                    if "220a" in nomeFonte or "220" in nomeFonte or "220 amperes" in nomeFonte or "220amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina 30A":       
                if "usina" in nomeFonte and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "smart" not in nomeFonte and "samrt" not in nomeFonte and "24v" in nomeFonte:
                    if "30a" in nomeFonte or "30" in nomeFonte or "30 amperes" in nomeFonte or "30amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina 70A":       
                if "usina" in nomeFonte and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "smart" not in nomeFonte and "samrt" not in nomeFonte and "24v" in nomeFonte:
                    if "70a" in nomeFonte or "70" in nomeFonte or "70 amperes" in nomeFonte or "70amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina 100A":       
                if "usina" in nomeFonte and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "smart" not in nomeFonte and "samrt" not in nomeFonte and "24v" in nomeFonte:
                    if "100a" in nomeFonte or "100" in nomeFonte or "100 amperes" in nomeFonte or "100amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Conversor de Tensao 30A":       
                if "usina" in nomeFonte and "conversor" in nomeFonte:
                    if "30a" in nomeFonte or "30" in nomeFonte or "30 amperes" in nomeFonte or "30amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Conversor de Tensao 60A":       
                if "usina" in nomeFonte and "conversor" in nomeFonte:
                    if "60a" in nomeFonte or "60" in nomeFonte or "60 amperes" in nomeFonte or "60amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Conversor de Tensao 120A":       
                if "usina" in nomeFonte and "conversor" in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Conversor de Tensao 240A":       
                if "usina" in nomeFonte and "conversor" in nomeFonte:
                    if "240a" in nomeFonte or "240" in nomeFonte or "240 amperes" in nomeFonte or "240amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Carregador de Baterias Charger 60A":       
                if "usina" in nomeFonte and "charger" in nomeFonte and "12v" in nomeFonte:
                    if "60a" in nomeFonte or "60" in nomeFonte or "60 amperes" in nomeFonte or "60amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)   
    
    def parse_renovonline(self, response):
        loja = "RENOV ONLINE"
        lugar = "São João da Boa Vista - SP"
        for i in response.xpath('//li[@class="ui-search-layout__item shops__layout-item shops__layout-item ui-search-layout__stack"]'):
            nomeFonte = i.xpath('.//div/div/div/div/a/h2/text()').get()
            price = i.xpath('.//div/div/div[3]/div/div[1]/div/div/div/div/span/span[2]/text()').get()
            cents = i.xpath('.//div/div/div[3]/div/div[1]/div/div/div/div/span/span[4]/text()').get()
            url = i.xpath('.//div/div/div[3]/div[1]/a/@href').get()
            nomeFonte = nomeFonte.lower()
            nomeFonte = unidecode.unidecode(nomeFonte)
            if not cents:
                cents = 0
            if price:
                price = price.replace('.', '')
                total_price = float(f"{price}.{cents}")
            if self.option_selected == "Fonte Usina Bob 60A":       
                if "bob" in nomeFonte and "usina" in nomeFonte and "smart" not in nomeFonte and "samrt" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "60a" in nomeFonte or "60" in nomeFonte or "60 amperes" in nomeFonte or "60amperes" in nomeFonte:
                            yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Bob 120A":       
                if "bob" in nomeFonte and "usina" in nomeFonte and "smart" not in nomeFonte and "samrt" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Bob 200A":       
                if "bob" in nomeFonte and "usina" in nomeFonte and "smart" not in nomeFonte and "samrt" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte and "12v" in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Battery Meter 50A":       
                if "usina" in nomeFonte and "battery" in nomeFonte and "meter" in nomeFonte and "bob" not in nomeFonte and "12v" in nomeFonte:
                    if "50a" in nomeFonte or "50" in nomeFonte or "50 amperes" in nomeFonte or "50amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Battery Meter 70A":       
                if "usina" in nomeFonte and "battery" in nomeFonte and "meter" in nomeFonte and "bob" not in nomeFonte and "12v" in nomeFonte:
                    if "50a" in nomeFonte or "50" in nomeFonte or "50 amperes" in nomeFonte or "50amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Battery Meter 100A":       
                if "usina" in nomeFonte and "battery" in nomeFonte and "meter" in nomeFonte and "bob" not in nomeFonte and "12v" in nomeFonte:
                    if "100a" in nomeFonte or "100" in nomeFonte or "100 amperes" in nomeFonte or "100amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Battery Meter 120A":       
                if "usina" in nomeFonte and "battery" in nomeFonte and "meter" in nomeFonte and "bob" not in nomeFonte and "12v" in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Smart 50A":       
                if "usina" in nomeFonte and ("smart" in nomeFonte or "samrt" in nomeFonte) and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "50a" in nomeFonte or "50" in nomeFonte or "50 amperes" in nomeFonte or "50amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Smart 70A":       
                if "usina" in nomeFonte and ("smart" in nomeFonte or "samrt" in nomeFonte) and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "70a" in nomeFonte or "70" in nomeFonte or "70 amperes" in nomeFonte or "70amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Smart 100A":       
                if "usina" in nomeFonte and ("smart" in nomeFonte or "samrt" in nomeFonte) and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "100a" in nomeFonte or "100" in nomeFonte or "100 amperes" in nomeFonte or "100amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Smart 120A":       
                if "usina" in nomeFonte and ("smart" in nomeFonte or "samrt" in nomeFonte) and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Smart 160A":       
                if "usina" in nomeFonte and ("smart" in nomeFonte or "samrt" in nomeFonte) and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "160a" in nomeFonte or "160" in nomeFonte or "160 amperes" in nomeFonte or "160amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Smart 200A MONO":       
                if "usina" in nomeFonte and ("smart" in nomeFonte or "samrt" in nomeFonte) and ("mono" in nomeFonte or "220v" in nomeFonte or "monovolt" in nomeFonte) and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina Smart 200A":       
                if "usina" in nomeFonte and ("samrt" in nomeFonte or "samrt" in nomeFonte) and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "12v" in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina 220A":       
                if "usina" in nomeFonte and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "smart" not in nomeFonte and "samrt" not in nomeFonte and "12v" in nomeFonte:
                    if "220a" in nomeFonte or "220" in nomeFonte or "220 amperes" in nomeFonte or "220amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina 30A":       
                if "usina" in nomeFonte and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "smart" not in nomeFonte and "samrt" not in nomeFonte and "24v" in nomeFonte:
                    if "30a" in nomeFonte or "30" in nomeFonte or "30 amperes" in nomeFonte or "30amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina 70A":       
                if "usina" in nomeFonte and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "smart" not in nomeFonte and "samrt" not in nomeFonte and "24v" in nomeFonte:
                    if "70a" in nomeFonte or "70" in nomeFonte or "70 amperes" in nomeFonte or "70amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Fonte Usina 100A":       
                if "usina" in nomeFonte and "bob" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte and "smart" not in nomeFonte and "samrt" not in nomeFonte and "24v" in nomeFonte:
                    if "100a" in nomeFonte or "100" in nomeFonte or "100 amperes" in nomeFonte or "100amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Conversor de Tensao 30A":       
                if "usina" in nomeFonte and "conversor" in nomeFonte:
                    if "30a" in nomeFonte or "30" in nomeFonte or "30 amperes" in nomeFonte or "30amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Conversor de Tensao 60A":       
                if "usina" in nomeFonte and "conversor" in nomeFonte:
                    if "60a" in nomeFonte or "60" in nomeFonte or "60 amperes" in nomeFonte or "60amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Conversor de Tensao 120A":       
                if "usina" in nomeFonte and "conversor" in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Conversor de Tensao 240A":       
                if "usina" in nomeFonte and "conversor" in nomeFonte:
                    if "240a" in nomeFonte or "240" in nomeFonte or "240 amperes" in nomeFonte or "240amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
            if self.option_selected == "Carregador de Baterias Charger 60A":       
                if "usina" in nomeFonte and "charger" in nomeFonte and "12v" in nomeFonte:
                    if "60a" in nomeFonte or "60" in nomeFonte or "60 amperes" in nomeFonte or "60amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

    def parse_shoppratico(self, response):
        loja = "SHOPPRATICO"
        lugar = "Sorocaba, São Paulo."
        for i in response.xpath('//li[@class="ui-search-layout__item shops__layout-item shops__layout-item ui-search-layout__stack"]'):
            nomeFonte = i.xpath('.//div/div/div/div/a/h2/text()').get()
            price = i.xpath('.//div/div/div[3]/div/div[1]/div/div/div/div/span/span[2]/text()').get()
            cents = i.xpath('.//div/div/div[3]/div/div[1]/div/div/div/div/span/span[4]/text()').get()
            url = i.xpath('.//div/div/div[3]/div[2]/a/@href').get()
            nomeFonte = nomeFonte.lower()
            nomeFonte = unidecode.unidecode(nomeFonte)
            if not cents:
                cents = 0
            if price:
                price = price.replace('.', '')
                total_price = float(f"{price}.{cents}")
            if self.option_selected == "Fonte Usina Bob 60A":       
                if "bob" in nomeFonte and "usina" in nomeFonte and "smart" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte:
                    if "60a" in nomeFonte or "60" in nomeFonte or "60 amperes" in nomeFonte or "60amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                        
            if self.option_selected == "Fonte Usina Bob 120A":       
                if "bob" in nomeFonte and "usina" in nomeFonte and "smart" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                        
            if self.option_selected == "Fonte Usina Bob 200A":       
                if "bob" in nomeFonte and "usina" in nomeFonte and "smart" not in nomeFonte and "battery" not in nomeFonte and "meter" not in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                        
            if self.option_selected == "Fonte Usina Battery Meter 50A":       
                if "usina" in nomeFonte and "battery" in nomeFonte and "meter" in nomeFonte and "smart" not in nomeFonte and "bob" not in nomeFonte:
                    if "50a" in nomeFonte or "50" in nomeFonte or "50 amperes" in nomeFonte or "50amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                        
            if self.option_selected == "Fonte Usina Battery Meter 70A":       
                if "usina" in nomeFonte and "battery" in nomeFonte and "meter" in nomeFonte and "smart" not in nomeFonte and "bob" not in nomeFonte:
                    if "50a" in nomeFonte or "50" in nomeFonte or "50 amperes" in nomeFonte or "50amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                        
            if self.option_selected == "Fonte Usina Battery Meter 100A":       
                if "usina" in nomeFonte and "battery" in nomeFonte and "meter" in nomeFonte and "smart" not in nomeFonte and "bob" not in nomeFonte:
                    if "100a" in nomeFonte or "100" in nomeFonte or "100 amperes" in nomeFonte or "100amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                        
            if self.option_selected == "Fonte Usina Battery Meter 120A":       
                if "usina" in nomeFonte and "battery" in nomeFonte and "meter" in nomeFonte and "smart" not in nomeFonte and "bob" not in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                        
            if self.option_selected == "Fonte Usina Smart 50A":       
                if "usina" in nomeFonte and "smart" in nomeFonte and "bob" not in nomeFonte and "battery" in nomeFonte and "meter" not in nomeFonte:
                    if "50a" in nomeFonte or "50" in nomeFonte or "50 amperes" in nomeFonte or "50amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                        
            if self.option_selected == "Fonte Usina Smart 70A":       
                if "usina" in nomeFonte and "smart" in nomeFonte and "bob" not in nomeFonte and "battery" in nomeFonte and "meter" not in nomeFonte:
                    if "70a" in nomeFonte or "70" in nomeFonte or "70 amperes" in nomeFonte or "70amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                        
            if self.option_selected == "Fonte Usina Smart 100A":       
                if "usina" in nomeFonte and "smart" in nomeFonte and "bob" not in nomeFonte and "battery" in nomeFonte and "meter" not in nomeFonte:
                    if "100a" in nomeFonte or "100" in nomeFonte or "100 amperes" in nomeFonte or "100amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                        
            if self.option_selected == "Fonte Usina Smart 120A":       
                if "usina" in nomeFonte and "smart" in nomeFonte and "bob" not in nomeFonte and "battery" in nomeFonte and "meter" not in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                        
            if self.option_selected == "Fonte Usina Smart 160A":       
                if "usina" in nomeFonte and "smart" in nomeFonte and "bob" not in nomeFonte and "battery" in nomeFonte and "meter" not in nomeFonte:
                    if "160a" in nomeFonte or "160" in nomeFonte or "160 amperes" in nomeFonte or "160amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                        
            if self.option_selected == "Fonte Usina Smart 200A MONO":       
                if "usina" in nomeFonte and "smart" in nomeFonte and ("mono" in nomeFonte or "220v" in nomeFonte or "monovolt" in nomeFonte) and "bob" not in nomeFonte and "battery" in nomeFonte and "meter" not in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                        
            if self.option_selected == "Fonte Usina Smart 200A":       
                if "usina" in nomeFonte and "smart" in nomeFonte and "bob" not in nomeFonte and "battery" in nomeFonte and "meter" not in nomeFonte:
                    if "200a" in nomeFonte or "200" in nomeFonte or "200 amperes" in nomeFonte or "200amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                        
            if self.option_selected == "Fonte Usina 220A":       
                if "usina" in nomeFonte and "bob" not in nomeFonte and "battery" in nomeFonte and "meter" not in nomeFonte and "smart" not in nomeFonte:
                    if "220a" in nomeFonte or "220" in nomeFonte or "220 amperes" in nomeFonte or "220amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                        
            if self.option_selected == "Fonte Usina 30A":       
                if "usina" in nomeFonte and "bob" not in nomeFonte and "battery" in nomeFonte and "meter" not in nomeFonte and "smart" not in nomeFonte:
                    if "30a" in nomeFonte or "30" in nomeFonte or "30 amperes" in nomeFonte or "30amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                        
            if self.option_selected == "Fonte Usina 70A":       
                if "usina" in nomeFonte and "bob" not in nomeFonte and "battery" in nomeFonte and "meter" not in nomeFonte and "smart" not in nomeFonte:
                    if "70a" in nomeFonte or "70" in nomeFonte or "70 amperes" in nomeFonte or "70amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                        
            if self.option_selected == "Fonte Usina 100A":       
                if "usina" in nomeFonte and "bob" not in nomeFonte and "battery" in nomeFonte and "meter" not in nomeFonte and "smart" not in nomeFonte:
                    if "100a" in nomeFonte or "100" in nomeFonte or "100 amperes" in nomeFonte or "100amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)     
        
            if self.option_selected == "Conversor de Tensao 30A":       
                if "usina" in nomeFonte and "conversor" in nomeFonte:
                    if "30a" in nomeFonte or "30" in nomeFonte or "30 amperes" in nomeFonte or "30amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                        
            if self.option_selected == "Conversor de Tensao 60A":       
                if "usina" in nomeFonte and "conversor" in nomeFonte:
                    if "60a" in nomeFonte or "60" in nomeFonte or "60 amperes" in nomeFonte or "60amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                                                    
            if self.option_selected == "Conversor de Tensao 120A":       
                if "usina" in nomeFonte and "conversor" in nomeFonte:
                    if "120a" in nomeFonte or "120" in nomeFonte or "120 amperes" in nomeFonte or "120amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

            if self.option_selected == "Conversor de Tensao 240A":       
                if "usina" in nomeFonte and "conversor" in nomeFonte:
                    if "240a" in nomeFonte or "240" in nomeFonte or "240 amperes" in nomeFonte or "240amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)
                                                
            if self.option_selected == "Carregador de Baterias Charger 60A":       
                if "usina" in nomeFonte and "charger" in nomeFonte:
                    if "60a" in nomeFonte or "60" in nomeFonte or "60 amperes" in nomeFonte or "60amperes" in nomeFonte:
                        yield from self.finish(total_price, url, nomeFonte, loja, lugar)

    def get_price_previsto(self, tipo):
        if tipo == "Clássico":
            if self.option_selected_new == "Fonte Usina Bob 60A":
                return fonteBob60Classico
            elif self.option_selected_new == "Fonte Usina Bob 120A":
                return fonteBob120Classico
            elif self.option_selected_new == "Fonte Usina Bob 200A":
                return fonteBob200Classico
            elif self.option_selected_new == "Fonte Usina Battery Meter 50A":
                return fonteBaterryMeter50Classico
            elif self.option_selected_new == "Fonte Usina Battery Meter 70A":
                return fonteBaterryMeter70Classico
            elif self.option_selected_new == "Fonte Usina Battery Meter 100A":
                return fonteBaterryMeter100Classico
            elif self.option_selected_new == "Fonte Usina Battery Meter 120A":
                return fonteBaterryMeter120Classico
            elif self.option_selected_new == "Fonte Usina Smart 50A":
                return fonteSmart50Classico
            elif self.option_selected_new == "Fonte Usina Smart 70A":
                return fonteSmart70Classico
            elif self.option_selected_new == "Fonte Usina Smart 100A":
                return fonteSmart100Classico
            elif self.option_selected_new == "Fonte Usina Smart 120A":
                return fonteSmart120Classico
            elif self.option_selected_new == "Fonte Usina Smart 160A":
                return fonteSmart160Classico
            elif self.option_selected_new == "Fonte Usina Smart 200A MONO":
                return fonteSmart200MonoClassico
            elif self.option_selected_new == "Fonte Usina Smart 200A":
                return fonteSmart200Classico
            elif self.option_selected_new == "Fonte Usina 220A":
                return fonteHeavyDuty220Classico
            elif self.option_selected_new == "Fonte Usina 30A":
                return fonte30Classico
            elif self.option_selected_new == "Fonte Usina 70A":
                return fonte70Classico
            elif self.option_selected_new == "Fonte Usina 100A":
                return fonte100Classico
            elif self.option_selected_new == "Conversor de Tensao 30A":
                return ConversorDeTensao30AClassico
            elif self.option_selected_new == "Conversor de Tensao 60A":
                return ConversorDeTensao60AClassico
            elif self.option_selected_new == "Conversor de Tensao 120A":
                return ConversorDeTensao120AClassico
            elif self.option_selected_new == "Conversor de Tensao 240A":
                return ConversorDeTensao240AClassico
            elif self.option_selected_new == "Carregador de Baterias Charger 60A":
                return CarregadorDeBateriasCharger60AClassico
        elif tipo == "Premium":
            if self.option_selected_new == "Fonte Usina Bob 60A":
                return fonteBob60Premium
            elif self.option_selected_new == "Fonte Usina Bob 120A":
                return fonteBob120Premium
            elif self.option_selected_new == "Fonte Usina Bob 200A":
                return fonteBob200Premium
            elif self.option_selected_new == "Fonte Usina Battery Meter 50A":
                return fonteBaterryMeter50Premium
            elif self.option_selected_new == "Fonte Usina Battery Meter 70A":
                return fonteBaterryMeter70Premium
            elif self.option_selected_new == "Fonte Usina Battery Meter 100A":
                return fonteBaterryMeter100Premium
            elif self.option_selected_new == "Fonte Usina Battery Meter 120A":
                return fonteBaterryMeter120Premium
            elif self.option_selected_new == "Fonte Usina Smart 50A":
                return fonteSmart50Premium
            elif self.option_selected_new == "Fonte Usina Smart 70A":
                return fonteSmart70Premium
            elif self.option_selected_new == "Fonte Usina Smart 100A":
                return fonteSmart100Premium
            elif self.option_selected_new == "Fonte Usina Smart 120A":
                return fonteSmart120Premium
            elif self.option_selected_new == "Fonte Usina Smart 160A":
                return fonteSmart160Premium
            elif self.option_selected_new == "Fonte Usina Smart 200A MONO":
                return fonteSmart200MonoPremium
            elif self.option_selected_new == "Fonte Usina Smart 200A":
                return fonteSmart200Premium
            elif self.option_selected_new == "Fonte Usina 220A":
                return fonteHeavyDuty220Premium
            elif self.option_selected_new == "Fonte Usina 30A":
                return fonte30Premium
            elif self.option_selected_new == "Fonte Usina 70A":
                return fonte70Premium
            elif self.option_selected_new == "Fonte Usina 100A":
                return fonte100Premium
            elif self.option_selected_new == "Conversor de Tensao 30A":
                return ConversorDeTensao30APremium
            elif self.option_selected_new == "Conversor de Tensao 60A":
                return ConversorDeTensao60APremium
            elif self.option_selected_new == "Conversor de Tensao 120A":
                return ConversorDeTensao120APremium
            elif self.option_selected_new == "Conversor de Tensao 240A":
                return ConversorDeTensao240APremium
            elif self.option_selected_new == "Carregador de Baterias Charger 60A":
                return CarregadorDeBateriasCharger60APremium
        elif tipo == "NA":
            if self.option_selected_new == "Fonte Usina Bob 60A":
                return fonteBob60Marketplace
            elif self.option_selected_new == "Fonte Usina Bob 120A":
                return fonteBob120Marketplace
            elif self.option_selected_new == "Fonte Usina Bob 200A":
                return fonteBob200Marketplace
            elif self.option_selected_new == "Fonte Usina Battery Meter 50A":
                return fonteBaterryMeter50Marketplace
            elif self.option_selected_new == "Fonte Usina Battery Meter 70A":
                return fonteBaterryMeter70Marketplace
            elif self.option_selected_new == "Fonte Usina Battery Meter 100A":
                return fonteBaterryMeter100Marketplace
            elif self.option_selected_new == "Fonte Usina Battery Meter 120A":
                return fonteBaterryMeter120Marketplace
            elif self.option_selected_new == "Fonte Usina Smart 50A":
                return fonteSmart50Marketplace
            elif self.option_selected_new == "Fonte Usina Smart 70A":
                return fonteSmart70Marketplace
            elif self.option_selected_new == "Fonte Usina Smart 100A":
                return fonteSmart100Marketplace
            elif self.option_selected_new == "Fonte Usina Smart 120A":
                return fonteSmart120Marketplace
            elif self.option_selected_new == "Fonte Usina Smart 160A":
                return fonteSmart160Marketplace
            elif self.option_selected_new == "Fonte Usina Smart 200A MONO":
                return fonteSmart200MonoMarketplace
            elif self.option_selected_new == "Fonte Usina Smart 200A":
                return fonteSmart200Marketplace
            elif self.option_selected_new == "Fonte Usina 220A":
                return fonteHeavyDuty220Marketplace
            elif self.option_selected_new == "Fonte Usina 30A":
                return fonte30Marketplace
            elif self.option_selected_new == "Fonte Usina 70A":
                return fonte70Marketplace
            elif self.option_selected_new == "Fonte Usina 100A":
                return fonte100Marketplace
            elif self.option_selected_new == "Conversor de Tensao 30A":
                return ConversorDeTensao30AMarketplace
            elif self.option_selected_new == "Conversor de Tensao 60A":
                return ConversorDeTensao60AMarketplace
            elif self.option_selected_new == "Conversor de Tensao 120A":
                return ConversorDeTensao120AMarketplace
            elif self.option_selected_new == "Conversor de Tensao 240A":
                return ConversorDeTensao240AMarketplace
            elif self.option_selected_new == "Carregador de Baterias Charger 60A":
                return CarregadorDeBateriasCharger60AMarketplace

    def parse_location(self, response):
        name = response.meta['name']
        url = response.meta['url']
        new_price_float = response.meta['price']
        tipo = response.meta['tipo']
        parcelado = self.get_price_previsto(tipo)
        loja = response.meta['loja']
        lugar = response.xpath('//*[@id="profile"]/div/div[2]/div[1]/div[3]/p/text()').get()


        doc.add_paragraph(f'Modelo: {self.option_selected_new}')
        doc.add_paragraph(f'URL: {url}')
        doc.add_paragraph(f'Nome: {name}')
        doc.add_paragraph(f'Preço: {new_price_float}')
        doc.add_paragraph(f'Preço Previsto: {parcelado}')
        doc.add_paragraph(f'Loja: {loja}')
        doc.add_paragraph(f'Tipo: {tipo}')
        doc.add_paragraph(f'Lugar: {lugar}')
        doc.add_paragraph("--------------------------------------------------------------------")
        doc.add_paragraph('')
        
        yield {
            'url': url,
            'name': name,
            'price': new_price_float,
            'price_previsto': parcelado,
            'loja': loja,
            'tipo': tipo,
            'lugar': lugar
        }
        doc.save(fr"dados/{self.option_selected_new}.docx")

        
        
