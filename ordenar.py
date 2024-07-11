import re
import docx
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH 
from docx.shared import Inches 
from spire.doc.common import *
from spire.doc import *


naoIndentificado= []
FonteUsinaBob60A = []
FonteUsinaBob200A = []
FonteUsinaBob120A = []
FonteUsinaBatteryMeter50A = []
FonteUsinaBatteryMeter70A = []
FonteUsinaBatteryMeter100A = []
FonteUsinaBatteryMeter120A = []
FonteUsinaSmart50A = []
FonteUsinaSmart70A = []
FonteUsinaSmart100A = []
FonteUsinaSmart120A = []
FonteUsinaSmart160A = []
FonteUsinaSmart200AMONO = []
FonteUsinaSmart200A = []
FonteUsina220A = []
FonteUsina30A = []
FonteUsina70A = []
FonteUsina100A = []
ConversorDeTensao30A = []
ConversorDeTensao60A = []
ConversorDeTensao120A = []
ConversorDeTensao240A =  []
CarregadorDeBateriasCharger60A = []

def read_docx(file_path):
    doc = docx.Document(file_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def read_text(text):
    items = []
    current_item = {}

    # Divide o texto em itens separados
    item_texts = re.split(r'-{5,}', text)

    for item_text in item_texts:
        lines = item_text.strip().split('\n')
        for line in lines:
            if line.startswith("Modelo:"):
                if current_item:
                    if current_item:
                        if current_item['Modelo'] == "Nao indentificado":
                            naoIndentificado.append(format_item_dif(current_item))
                        if current_item['Modelo'] == "Fonte Usina Bob 60A":
                            FonteUsinaBob60A.append((format_item(current_item), current_item['Loja']))
                        if current_item['Modelo'] == "Fonte Usina Bob 120A":
                            FonteUsinaBob120A.append((format_item(current_item), current_item['Loja']))
                        if current_item['Modelo'] == "Fonte Usina Bob 200A":
                            FonteUsinaBob200A.append((format_item(current_item), current_item['Loja']))
                        if current_item['Modelo'] == "Fonte Usina Battery Meter 50A":
                            FonteUsinaBatteryMeter50A.append((format_item(current_item), current_item['Loja']))
                        if current_item['Modelo'] == "Fonte Usina Battery Meter 70A":
                            FonteUsinaBatteryMeter70A.append((format_item(current_item), current_item['Loja']))
                        if current_item['Modelo'] == "Fonte Usina Battery Meter 100A":
                            FonteUsinaBatteryMeter100A.append((format_item(current_item), current_item['Loja']))
                        if current_item['Modelo'] == "Fonte Usina Battery Meter 120A":
                            FonteUsinaBatteryMeter120A.append((format_item(current_item), current_item['Loja']))
                        if current_item['Modelo'] == "Fonte Usina Smart 50A":
                            FonteUsinaSmart50A.append((format_item(current_item), current_item['Loja']))
                        if current_item['Modelo'] == "Fonte Usina Smart 70A":
                            FonteUsinaSmart70A.append((format_item(current_item), current_item['Loja']))
                        if current_item['Modelo'] == "Fonte Usina Smart 100A":
                            FonteUsinaSmart100A.append((format_item(current_item), current_item['Loja']))
                        if current_item['Modelo'] == "Fonte Usina Smart 120A":
                            FonteUsinaSmart120A.append((format_item(current_item), current_item['Loja']))
                        if current_item['Modelo'] == "Fonte Usina Smart 160A":
                            FonteUsinaSmart160A.append((format_item(current_item), current_item['Loja']))
                        if current_item['Modelo'] == "Fonte Usina Smart 200A MONO":
                            FonteUsinaSmart200AMONO.append((format_item(current_item), current_item['Loja']))
                        if current_item['Modelo'] == "Fonte Usina Smart 200A":
                            FonteUsinaSmart200A.append((format_item(current_item), current_item['Loja']))
                        if current_item['Modelo'] == "Fonte Usina 220A":
                            FonteUsina220A.append((format_item(current_item), current_item['Loja']))
                        if current_item['Modelo'] == "Fonte Usina 30A":
                            FonteUsina30A.append((format_item(current_item), current_item['Loja']))
                        if current_item['Modelo'] == "Fonte Usina 70A":
                            FonteUsina70A.append((format_item(current_item), current_item['Loja']))
                        if current_item['Modelo'] == "Fonte Usina 100A":
                            FonteUsina100A.append((format_item(current_item), current_item['Loja']))
                        if current_item['Modelo'] == "Conversor de Tensao 30A":
                            ConversorDeTensao30A.append((format_item(current_item), current_item['Loja']))
                        if current_item['Modelo'] == "Conversor de Tensao 60A":
                            ConversorDeTensao60A.append((format_item(current_item), current_item['Loja']))
                        if current_item['Modelo'] == "Conversor de Tensao 120A":
                            ConversorDeTensao120A.append((format_item(current_item), current_item['Loja']))
                        if current_item['Modelo'] == "Conversor de Tensao 240A":
                            ConversorDeTensao240A.append((format_item(current_item), current_item['Loja']))
                        if current_item['Modelo'] == "Carregador de Baterias Charger 60A":
                            CarregadorDeBateriasCharger60A.append((format_item(current_item), current_item['Loja']))
                        current_item = {}
                current_item['Modelo'] = line.split("Modelo:", 1)[1].strip()
                
            elif line.startswith("URL:"):
                current_item['URL'] = line.split("URL:", 1)[1].strip()
            elif line.startswith("Nome:"):
                current_item['Nome'] = line.split("Nome:", 1)[1].strip()
            elif line.startswith("Preço:"):
                current_item['Preço'] = line.split("Preço:", 1)[1].strip()
            elif line.startswith("Preço Previsto:"):
                current_item['Preço Previsto'] = line.split("Preço Previsto:", 1)[1].strip()
            elif line.startswith("Loja:"):
                current_item['Loja'] = line.split("Loja:", 1)[1].strip()
            elif line.startswith("Tipo:"):
                current_item['Tipo'] = line.split("Tipo:", 1)[1].strip()
            elif line.startswith("Lugar:"):
                current_item['Lugar'] = line.split("Lugar:", 1)[1].strip()
                
        if current_item:
            if current_item['Modelo'] == "Nao indentificado":
                naoIndentificado.append(format_item_dif(current_item))
            if current_item['Modelo'] == "Fonte Usina Bob 60A":
                FonteUsinaBob60A.append((format_item(current_item), current_item['Loja']))
            if current_item['Modelo'] == "Fonte Usina Bob 120A":
                FonteUsinaBob120A.append((format_item(current_item), current_item['Loja']))
            if current_item['Modelo'] == "Fonte Usina Bob 200A":
                FonteUsinaBob200A.append((format_item(current_item), current_item['Loja']))
            if current_item['Modelo'] == "Fonte Usina Battery Meter 50A":
                FonteUsinaBatteryMeter50A.append((format_item(current_item), current_item['Loja']))
            if current_item['Modelo'] == "Fonte Usina Battery Meter 70A":
                FonteUsinaBatteryMeter70A.append((format_item(current_item), current_item['Loja']))
            if current_item['Modelo'] == "Fonte Usina Battery Meter 100A":
                FonteUsinaBatteryMeter100A.append((format_item(current_item), current_item['Loja']))
            if current_item['Modelo'] == "Fonte Usina Battery Meter 120A":
                FonteUsinaBatteryMeter120A.append((format_item(current_item), current_item['Loja']))
            if current_item['Modelo'] == "Fonte Usina Smart 50A":
                FonteUsinaSmart50A.append((format_item(current_item), current_item['Loja']))
            if current_item['Modelo'] == "Fonte Usina Smart 70A":
                FonteUsinaSmart70A.append((format_item(current_item), current_item['Loja']))
            if current_item['Modelo'] == "Fonte Usina Smart 100A":
                FonteUsinaSmart100A.append((format_item(current_item), current_item['Loja']))
            if current_item['Modelo'] == "Fonte Usina Smart 120A":
                FonteUsinaSmart120A.append((format_item(current_item), current_item['Loja']))
            if current_item['Modelo'] == "Fonte Usina Smart 160A":
                FonteUsinaSmart160A.append((format_item(current_item), current_item['Loja']))
            if current_item['Modelo'] == "Fonte Usina Smart 200A MONO":
                FonteUsinaSmart200AMONO.append((format_item(current_item), current_item['Loja']))
            if current_item['Modelo'] == "Fonte Usina Smart 200A":
                FonteUsinaSmart200A.append((format_item(current_item), current_item['Loja']))
            if current_item['Modelo'] == "Fonte Usina 220A":
                FonteUsina220A.append((format_item(current_item), current_item['Loja']))
            if current_item['Modelo'] == "Fonte Usina 30A":
                FonteUsina30A.append((format_item(current_item), current_item['Loja']))
            if current_item['Modelo'] == "Fonte Usina 70A":
                FonteUsina70A.append((format_item(current_item), current_item['Loja']))
            if current_item['Modelo'] == "Fonte Usina 100A":
                FonteUsina100A.append((format_item(current_item), current_item['Loja']))
            if current_item['Modelo'] == "Conversor de Tensao 30A":
                ConversorDeTensao30A.append((format_item(current_item), current_item['Loja']))
            if current_item['Modelo'] == "Conversor de Tensao 60A":
                ConversorDeTensao60A.append((format_item(current_item), current_item['Loja']))
            if current_item['Modelo'] == "Conversor de Tensao 120A":
                ConversorDeTensao120A.append((format_item(current_item), current_item['Loja']))
            if current_item['Modelo'] == "Conversor de Tensao 240A":
                ConversorDeTensao240A.append((format_item(current_item), current_item['Loja']))
            if current_item['Modelo'] == "Carregador de Baterias Charger 60A":
                CarregadorDeBateriasCharger60A.append((format_item(current_item), current_item['Loja']))
            current_item = {}


def format_item(item):
    formatted_item = f"{item['Loja']} – {item['Lugar']} – Preço Anúncio: R$ {item['Preço']} – Preço Política: R$ {item['Preço Previsto']} ({item['Tipo']})\n{item['URL']}\n"
    return formatted_item

def format_item_dif(item):
    formatted_item = f"{item['URL']}\n"
    return formatted_item
lojas = {}

for item_path in os.listdir(r"dados/"):
    file_path = os.path.join(r"dados/", item_path)
    text = read_docx(file_path)

    read_text(text)

output_doc = docx.Document()
for item in FonteUsinaBob60A:
    if item[1] not in lojas:
        lojas[item[1]] = []
    lojas[item[1]].append((item[0], "Fonte Usina Bob 60A"))
for item in FonteUsinaBob120A:
    if item[1] not in lojas:
        lojas[item[1]] = []
    lojas[item[1]].append((item[0], "Fonte Usina Bob 120A"))
for item in FonteUsinaBob200A:
    if item[1] not in lojas:
        lojas[item[1]] = []
    lojas[item[1]].append((item[0], "Fonte Usina Bob 200A"))
for item in FonteUsinaBatteryMeter50A:
    if item[1] not in lojas:
        lojas[item[1]] = []
    lojas[item[1]].append((item[0], "Fonte Usina Battery Meter 50A"))
for item in FonteUsinaBatteryMeter70A:
    if item[1] not in lojas:
        lojas[item[1]] = []
    lojas[item[1]].append((item[0], "Fonte Usina Battery Meter 70A"))
for item in FonteUsinaBatteryMeter100A:
    if item[1] not in lojas:
        lojas[item[1]] = []
    lojas[item[1]].append((item[0], "Fonte Usina Battery Meter 100A"))
for item in FonteUsinaBatteryMeter120A:
    if item[1] not in lojas:
        lojas[item[1]] = []
    lojas[item[1]].append((item[0], "Fonte Usina Battery Meter 120A"))
for item in FonteUsinaSmart50A:
    if item[1] not in lojas:
        lojas[item[1]] = []
    lojas[item[1]].append((item[0], "Fonte Usina Smart 50A"))
for item in FonteUsinaSmart70A:
    if item[1] not in lojas:
        lojas[item[1]] = []
    lojas[item[1]].append((item[0], "Fonte Usina Smart 70A"))
for item in FonteUsinaSmart100A:
    if item[1] not in lojas:
        lojas[item[1]] = []
    lojas[item[1]].append((item[0], "Fonte Usina Smart 100A"))
for item in FonteUsinaSmart120A:
    if item[1] not in lojas:
        lojas[item[1]] = []
    lojas[item[1]].append((item[0], "Fonte Usina Smart 120A"))
for item in FonteUsinaSmart160A:
    if item[1] not in lojas:
        lojas[item[1]] = []
    lojas[item[1]].append((item[0], "Fonte Usina Smart 160A"))
for item in FonteUsinaSmart200AMONO:
    if item[1] not in lojas:
        lojas[item[1]] = []
    lojas[item[1]].append((item[0], "Fonte Usina Smart 200A MONO"))
for item in FonteUsinaSmart200A:
    if item[1] not in lojas:
        lojas[item[1]] = []
    lojas[item[1]].append((item[0], "Fonte Usina Smart 200A"))
for item in FonteUsina220A:
    if item[1] not in lojas:
        lojas[item[1]] = []
    lojas[item[1]].append((item[0], "Fonte Usina 220A"))
for item in FonteUsina30A:
    if item[1] not in lojas:
        lojas[item[1]] = []
    lojas[item[1]].append((item[0], "Fonte Usina 30A"))
for item in FonteUsina70A:
    if item[1] not in lojas:
        lojas[item[1]] = []
    lojas[item[1]].append((item[0], "Fonte Usina 70A"))
for item in FonteUsina100A:
    if item[1] not in lojas:
        lojas[item[1]] = []
    lojas[item[1]].append((item[0], "Fonte Usina 100A"))
for item in ConversorDeTensao30A:
    if item[1] not in lojas:
        lojas[item[1]] = []
    lojas[item[1]].append((item[0], "Conversor de Tensao 30A"))
for item in ConversorDeTensao60A:
    if item[1] not in lojas:
        lojas[item[1]] = []
    lojas[item[1]].append((item[0], "Conversor de Tensao 60A"))
for item in ConversorDeTensao120A:
    if item[1] not in lojas:
        lojas[item[1]] = []
    lojas[item[1]].append((item[0], "Conversor de Tensao 120A"))
for item in ConversorDeTensao240A:
    if item[1] not in lojas:
        lojas[item[1]] = []
    lojas[item[1]].append((item[0], "Conversor de Tensao 240A"))
for item in CarregadorDeBateriasCharger60A:
    if item[1] not in lojas:
        lojas[item[1]] = []
    lojas[item[1]].append((item[0], "Carregador de Baterias Charger 60A"))
for i in lojas:
    output_doc.add_paragraph().add_run(f"*{i}*\n").bold = True
    for item, modelo in lojas[i]:
        output_doc.add_paragraph(f"{modelo} - {item}").paragraph_format.left_indent = Inches(0.5)
        
output_doc.save(r'dados_extraidos.docx')