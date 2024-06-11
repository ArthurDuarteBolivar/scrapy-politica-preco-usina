import re
import docx
import os

naoIndentificado= []
storm40 = []
storm60 = []
lite60 = []
lite70 = []
storm70 = []
bob90 = []
storm120 = []
lite120 = []
bob120 = []
storm200 = []
lite200 = []
mono200 = []
bob200 = []


def read_docx(file_path):
    doc = docx.Document(file_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def read_text(text):
    items = []
    current_item = {}

    # Divide o texto em itens separados
    item_texts = re.split(r'-{5,}', text)

    for item_text in item_texts:
        lines = item_text.strip().split('\n')
        for line in lines:
            if line.startswith("Modelo:"):
                if current_item:
                    if current_item:
                        if current_item['Modelo'] == "Nao indentificado":
                            naoIndentificado.append(format_item_dif(current_item))
                        if current_item['Modelo'] == "Storm 40":
                            storm40.append(format_item(current_item))
                        if current_item['Modelo'] == "Lite 60":
                            lite60.append(format_item(current_item))
                        if current_item['Modelo'] == "Storm 60":
                            storm60.append(format_item(current_item))
                        if current_item['Modelo'] == "Lite 70":
                            lite70.append(format_item(current_item))
                        if current_item['Modelo'] == "Storm 70":
                            storm70.append(format_item(current_item))
                        if current_item['Modelo'] == "Bob 90":
                            bob90.append(format_item(current_item))
                        if current_item['Modelo'] == "Storm 120":
                            storm120.append(format_item(current_item))
                        if current_item['Modelo'] == "Lite 120":
                            lite120.append(format_item(current_item))
                        if current_item['Modelo'] == "Bob 120":
                            bob120.append(format_item(current_item))
                        if current_item['Modelo'] == "Storm 200":
                            storm200.append(format_item(current_item))
                        if current_item['Modelo'] == "Lite 200":
                            lite200.append(format_item(current_item))
                        if current_item['Modelo'] == "Bob 200":
                            bob200.append(format_item(current_item))
                        if current_item['Modelo'] == "Storm 200 MONO":
                            mono200.append(format_item(current_item))
                        # items.append(format_item(current_item))
                        current_item = {}
                current_item['Modelo'] = line.split("Modelo:", 1)[1].strip()
            elif line.startswith("URL:"):
                current_item['URL'] = line.split("URL:", 1)[1].strip()
            elif line.startswith("Nome:"):
                current_item['Nome'] = line.split("Nome:", 1)[1].strip()
            elif line.startswith("Preço:"):
                current_item['Preço'] = line.split("Preço:", 1)[1].strip()
            elif line.startswith("Preço Previsto:"):
                current_item['Preço Previsto'] = line.split("Preço Previsto:", 1)[1].strip()
            elif line.startswith("Loja:"):
                current_item['Loja'] = line.split("Loja:", 1)[1].strip()
            elif line.startswith("Tipo:"):
                current_item['Tipo'] = line.split("Tipo:", 1)[1].strip()
            elif line.startswith("Lugar:"):
                current_item['Lugar'] = line.split("Lugar:", 1)[1].strip()
        if current_item:
            if current_item['Modelo'] == "Nao indentificado":
                naoIndentificado.append(format_item_dif(current_item))
            if current_item['Modelo'] == "Storm 40":
                storm40.append(format_item(current_item))
            if current_item['Modelo'] == "Lite 60":
                lite60.append(format_item(current_item))
            if current_item['Modelo'] == "Storm 60":
                storm60.append(format_item(current_item))
            if current_item['Modelo'] == "Lite 70":
                lite70.append(format_item(current_item))
            if current_item['Modelo'] == "Storm 70":
                storm70.append(format_item(current_item))
            if current_item['Modelo'] == "Bob 90":
                bob90.append(format_item(current_item))
            if current_item['Modelo'] == "Storm 120":
                storm120.append(format_item(current_item))
            if current_item['Modelo'] == "Lite 120":
                lite120.append(format_item(current_item))
            if current_item['Modelo'] == "Bob 120":
                bob120.append(format_item(current_item))
            if current_item['Modelo'] == "Storm 200":
                storm200.append(format_item(current_item))
            if current_item['Modelo'] == "Lite 200":
                lite200.append(format_item(current_item))
            if current_item['Modelo'] == "Bob 200":
                bob200.append(format_item(current_item))
            if current_item['Modelo'] == "Storm 200 MONO":
                mono200.append(format_item(current_item))
            # items.append(format_item(current_item))
            current_item = {}


def format_item(item):
    formatted_item = f"{item['Loja']} – {item['Lugar']} – Preço Anúncio: R$ {item['Preço']} – Preço Política: R$ {item['Preço Previsto']} ({item['Tipo']})\n{item['URL']}\n"
    return formatted_item

def format_item_dif(item):
    formatted_item = f"{item['URL']}\n"
    return formatted_item

for item_path in os.listdir(r"C:\workspace\mercado-livre\mercadolivre\dados"):
    file_path = os.path.join(r"C:\workspace\mercado-livre\mercadolivre\dados", item_path)
    text = read_docx(file_path)

    read_text(text)

    output_doc = docx.Document()


    if naoIndentificado:
        output_doc.add_paragraph("*Modelo não indentificado*\n")
        for item in naoIndentificado:
            output_doc.add_paragraph(item)
    if storm40:
        output_doc.add_paragraph("*Storm 40*\n")
        for item in storm40:
            output_doc.add_paragraph(item)
    if lite60:
        output_doc.add_paragraph("*Lite 60*\n")
        for item in lite60:
            output_doc.add_paragraph(item)
    if storm60:
        output_doc.add_paragraph("*Storm 60*\n")
        for item in storm60:
            output_doc.add_paragraph(item)
    if lite70:
        output_doc.add_paragraph("*Lite 70*\n")
        for item in lite70:
            output_doc.add_paragraph(item)
    if storm70:
        output_doc.add_paragraph("*Storm 70*\n")
        for item in storm70:
            output_doc.add_paragraph(item)
    if bob90:
        output_doc.add_paragraph("*Bob 90*\n")
        for item in bob90:
            output_doc.add_paragraph(item)
    if bob120:
        output_doc.add_paragraph("*Bob 120*\n")
        for item in bob120:
            output_doc.add_paragraph(item)
    if lite120:
        output_doc.add_paragraph("*Lite 120*\n")
        for item in lite120:
            output_doc.add_paragraph(item)
    if storm120:
        output_doc.add_paragraph("*Storm 120*\n")
        for item in storm120:
            output_doc.add_paragraph(item)
    if bob200:
        output_doc.add_paragraph("*Bob 200*\n")
        for item in bob200:
            output_doc.add_paragraph(item)
    if lite200:
        output_doc.add_paragraph("*Lite 200*\n")
        for item in lite200:
            output_doc.add_paragraph(item)
    if mono200:
        output_doc.add_paragraph("*Storm 200 MONO*\n")
        for item in mono200:
            output_doc.add_paragraph(item)
    if storm200:
        output_doc.add_paragraph("*Storm 200*\n")
        for item in storm200:
            output_doc.add_paragraph(item)


    output_doc.save(r'C:\workspace\mercado-livre\mercadolivre\dados_extraidos.docx')